import asyncio
import logging
from datetime import datetime, timedelta, timezone
import os
import json
import re
import time
import sys
import signal
import hashlib
import hmac
from collections import defaultdict
from ipaddress import ip_address
from typing import Dict, List, Optional, Tuple
import html
import random

from PIL import Image
import pandas as pd
import openpyxl
import aiohttp
from bs4 import BeautifulSoup
import requests
from urllib.parse import quote_plus, urlparse, urljoin
import xml.etree.ElementTree as ET

# Для перевода
try:
    from googletrans import Translator, LANGUAGES
    import langdetect
    TRANSLATION_AVAILABLE = True
except ImportError:
    TRANSLATION_AVAILABLE = False
    print("⚠️ Библиотеки для перевода не установлены. Установите: pip install googletrans==4.0.0-rc1 langdetect")

from aiogram import Bot, Dispatcher, types
from aiogram.contrib.middlewares.logging import LoggingMiddleware
from aiogram.types import ParseMode, ContentType, InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardMarkup, KeyboardButton
from aiogram.utils import executor
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from telethon import TelegramClient
from telethon.errors import FloodWaitError, SessionPasswordNeededError, PhoneCodeExpiredError, PhoneCodeInvalidError
from telethon.sessions import StringSession
from telethon.tl.functions.messages import SearchGlobalRequest
from telethon.tl.types import InputMessagesFilterEmpty, InputMessagesFilterPhotos, InputMessagesFilterDocument, InputMessagesFilterVideo
from telethon.tl.types import PeerUser, PeerChat, PeerChannel

# ========== НАСТРОЙКИ БЕЗОПАСНОСТИ ==========
# В продакшене используй переменные окружения!
API_TOKEN = os.environ.get('BOT_TOKEN', '8029293386:AAG-Hih0eVun77YYcY8zf6PyDjQERmQCx9w')
API_ID = int(os.environ.get('API_ID', '38892524'))
API_HASH = os.environ.get('API_HASH', 'd71ef1a657ab20d2a47a52130626c939')
ADMIN_ID = int(os.environ.get('ADMIN_ID', '5224743551'))
SECRET_KEY = os.environ.get('SECRET_KEY', 'default-secret-key-change-this-in-production')

# Файлы для хранения данных
CHANNELS_FILE = 'saved_channels.json'
SESSIONS_DIR = 'telegram_sessions'
IMAGES_DIR = 'temp_images'
UPLOADS_DIR = 'uploads'
SECURITY_LOG_FILE = 'security_logs.json'
BLOCKED_IPS_FILE = 'blocked_ips.json'

# Часовой пояс UTC+10
UTC_PLUS_10 = timezone(timedelta(hours=10))

# Максимальный период в днях
MAX_PERIOD_DAYS = 30
MAX_PERIOD_HOURS = MAX_PERIOD_DAYS * 24

# Настройки для поиска в интернете
MAX_WEB_RESULTS = 100
SEARCH_ENGINES = ['google', 'bing', 'yahoo', 'yandex', 'duckduckgo', 'baidu', 'ask']
MAX_SEARCH_TIME = 120

# Настройки безопасности
SECURITY_SETTINGS = {
    'rate_limiting': True,
    'max_requests_per_minute': 60,
    'brute_force_protection': True,
    'max_auth_attempts': 5,
    'max_code_attempts': 3,
    'lockout_minutes': 5,
    'input_validation': True,
    'max_text_length': 5000,
    'sql_injection_protection': True,
    'xss_protection': True,
    'validate_telegram_ua': True,
    'log_suspicious': True,
    'log_failed_attempts': True,
    'notify_admin_on_attack': True
}

# Создаем папки
os.makedirs(SESSIONS_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)

# ========== ПРЕДОТВРАЩЕНИЕ КОНФЛИКТОВ ==========
def signal_handler(sig, frame):
    print('Остановка бота...')
    sys.exit(0)

signal.signal(signal.SIGINT, signal_handler)
signal.signal(signal.SIGTERM, signal_handler)

# ========== ИНИЦИАЛИЗАЦИЯ ==========
bot = Bot(token=API_TOKEN)
dp = Dispatcher(bot)
logging.basicConfig(level=logging.INFO)

# Временное хранилище для данных пользователя
user_data = {}
auth_data = {}
stop_flags = {}

# Глобальная сессия
MASTER_SESSION = None

# ========== СИСТЕМА БЕЗОПАСНОСТИ ==========
class SecurityManager:
    """Менеджер безопасности для защиты бота"""
    
    def __init__(self):
        self.request_counts = defaultdict(list)
        self.auth_attempts = defaultdict(list)
        self.code_attempts = defaultdict(list)
        self.blocked_users = {}
        self.blocked_ips = self._load_blocked_ips()
        self.security_logs = self._load_security_logs()
        
    def _load_blocked_ips(self):
        if os.path.exists(BLOCKED_IPS_FILE):
            with open(BLOCKED_IPS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def _save_blocked_ips(self):
        with open(BLOCKED_IPS_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.blocked_ips, f, ensure_ascii=False, indent=2)
    
    def _load_security_logs(self):
        if os.path.exists(SECURITY_LOG_FILE):
            with open(SECURITY_LOG_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return []
    
    def _save_security_logs(self):
        with open(SECURITY_LOG_FILE, 'w', encoding='utf-8') as f:
            json.dump(self.security_logs[-1000:], f, ensure_ascii=False, indent=2)
    
    def log_security_event(self, event_type: str, user_id: int, details: str, ip: str = None):
        if not SECURITY_SETTINGS['log_suspicious']:
            return
            
        log_entry = {
            'timestamp': datetime.now().isoformat(),
            'event_type': event_type,
            'user_id': user_id,
            'ip': ip or 'unknown',
            'details': details
        }
        self.security_logs.append(log_entry)
        self._save_security_logs()
        
        if SECURITY_SETTINGS['notify_admin_on_attack'] and 'attack' in event_type.lower():
            asyncio.create_task(self._notify_admin(log_entry))
    
    async def _notify_admin(self, log_entry: dict):
        try:
            await bot.send_message(
                ADMIN_ID,
                f"🚨 *Обнаружена атака!*\n\n"
                f"Тип: {log_entry['event_type']}\n"
                f"Пользователь: {log_entry['user_id']}\n"
                f"IP: {log_entry['ip']}\n"
                f"Время: {log_entry['timestamp']}\n"
                f"Детали: {log_entry['details']}",
                parse_mode='Markdown'
            )
        except:
            pass
    
    def check_rate_limit(self, user_id: int) -> bool:
        if not SECURITY_SETTINGS['rate_limiting']:
            return True
            
        now = time.time()
        minute_ago = now - 60
        
        self.request_counts[user_id] = [t for t in self.request_counts[user_id] if t > minute_ago]
        
        if len(self.request_counts[user_id]) >= SECURITY_SETTINGS['max_requests_per_minute']:
            self.log_security_event('rate_limit_exceeded', user_id, f'Превышен лимит запросов')
            return False
            
        self.request_counts[user_id].append(now)
        return True
    
    def check_brute_force(self, user_id: int, attempt_type: str = 'auth') -> bool:
        if not SECURITY_SETTINGS['brute_force_protection']:
            return True
            
        now = time.time()
        
        if user_id in self.blocked_users:
            if now < self.blocked_users[user_id]:
                return False
            else:
                del self.blocked_users[user_id]
        
        attempts_list = self.auth_attempts if attempt_type == 'auth' else self.code_attempts
        max_attempts = SECURITY_SETTINGS['max_auth_attempts'] if attempt_type == 'auth' else SECURITY_SETTINGS['max_code_attempts']
        
        attempts = attempts_list[user_id]
        attempts = [t for t in attempts if t > now - 1800]
        
        if len(attempts) >= max_attempts:
            lockout_time = now + (SECURITY_SETTINGS['lockout_minutes'] * 60)
            self.blocked_users[user_id] = lockout_time
            self.log_security_event(f'brute_force_attack_{attempt_type}', user_id, f'Заблокирован')
            return False
        
        attempts.append(now)
        attempts_list[user_id] = attempts
        return True
    
    def validate_input(self, text: str, max_length: int = None) -> Tuple[bool, str]:
        if not SECURITY_SETTINGS['input_validation']:
            return True, text
            
        max_len = max_length or SECURITY_SETTINGS['max_text_length']
        if len(text) > max_len:
            return False, f"Текст слишком длинный (макс. {max_len} символов)"
        
        if SECURITY_SETTINGS['sql_injection_protection']:
            sql_patterns = [
                r'\bSELECT\b.*\bFROM\b', r'\bINSERT\b.*\bINTO\b', r'\bUPDATE\b.*\bSET\b',
                r'\bDELETE\b.*\bFROM\b', r'\bDROP\b.*\bTABLE\b', r'\bUNION\b.*\bSELECT\b',
                r'--', r';'
            ]
            for pattern in sql_patterns:
                if re.search(pattern, text, re.IGNORECASE):
                    self.log_security_event('sql_injection_attempt', 0, 'Обнаружена SQL-инъекция')
                    return False, "Обнаружена потенциальная SQL-инъекция"
        
        if SECURITY_SETTINGS['xss_protection']:
            xss_patterns = [
                r'<script.*?>.*?</script>', r'javascript:', r'on\w+\s*=',
                r'<.*?on\w+.*?>', r'<iframe.*?>', r'<object.*?>'
            ]
            for pattern in xss_patterns:
                if re.search(pattern, text, re.IGNORECASE | re.DOTALL):
                    self.log_security_event('xss_attempt', 0, 'Обнаружена XSS-атака')
                    return False, "Обнаружена потенциальная XSS-атака"
            
            text = html.escape(text)
        
        return True, text
    
    def validate_telegram_request(self, message: types.Message) -> bool:
        if not SECURITY_SETTINGS['validate_telegram_ua']:
            return True
            
        if not hasattr(message, 'from_user') or not hasattr(message, 'chat'):
            self.log_security_event('invalid_telegram_structure', 0, 'Неверная структура')
            return False
        
        if not message.from_user or not message.from_user.id:
            self.log_security_event('missing_user_id', 0, 'Отсутствует ID')
            return False
        
        return True
    
    def verify_owner(self, user_id: int) -> bool:
        return user_id == ADMIN_ID
    
    def get_security_report(self) -> dict:
        now = time.time()
        hour_ago = now - 3600
        
        return {
            'total_logs': len(self.security_logs),
            'logs_last_hour': len([l for l in self.security_logs 
                                  if datetime.fromisoformat(l['timestamp']).timestamp() > hour_ago]),
            'blocked_users': len(self.blocked_users),
            'blocked_ips': len(self.blocked_ips),
            'rate_limit_violations': len([l for l in self.security_logs 
                                         if l['event_type'] == 'rate_limit_exceeded']),
            'brute_force_attempts': len([l for l in self.security_logs 
                                        if 'brute_force' in l['event_type']])
        }

security = SecurityManager()

# ========== ФУНКЦИИ ДЛЯ ПЕРЕВОДА ==========
class TranslationManager:
    """Менеджер для перевода текстов"""
    
    def __init__(self):
        self.translator = Translator() if TRANSLATION_AVAILABLE else None
        self.supported_languages = LANGUAGES if TRANSLATION_AVAILABLE else {}
        
    def detect_language(self, text: str) -> str:
        """Определяет язык текста"""
        if not TRANSLATION_AVAILABLE or not text:
            return 'unknown'
        try:
            lang = langdetect.detect(text)
            return lang
        except:
            return 'unknown'
    
    def get_language_name(self, lang_code: str) -> str:
        """Возвращает название языка по коду"""
        if not TRANSLATION_AVAILABLE:
            return lang_code
        return self.supported_languages.get(lang_code, lang_code)
    
    async def translate(self, text: str, dest_lang: str = 'ru', src_lang: str = 'auto') -> dict:
        """Переводит текст с указанием источника"""
        if not TRANSLATION_AVAILABLE or not text:
            return {
                'original': text,
                'translated': text,
                'src_lang': 'unknown',
                'dest_lang': dest_lang,
                'was_translated': False
            }
        
        try:
            if src_lang == 'auto':
                src_lang = self.detect_language(text)
            
            if src_lang == dest_lang:
                return {
                    'original': text,
                    'translated': text,
                    'src_lang': src_lang,
                    'dest_lang': dest_lang,
                    'was_translated': False
                }
            
            result = await self.translator.translate(text, dest=dest_lang, src=src_lang)
            
            return {
                'original': text,
                'translated': result.text,
                'src_lang': src_lang,
                'dest_lang': dest_lang,
                'src_lang_name': self.get_language_name(src_lang),
                'dest_lang_name': self.get_language_name(dest_lang),
                'was_translated': True
            }
        except Exception as e:
            print(f"Ошибка перевода: {e}")
            return {
                'original': text,
                'translated': text,
                'src_lang': 'unknown',
                'dest_lang': dest_lang,
                'was_translated': False
            }
    
    def add_translation_note(self, text: str, src_lang: str) -> str:
        """Добавляет пометку о переводе"""
        if not TRANSLATION_AVAILABLE or src_lang == 'ru' or src_lang == 'unknown':
            return text
        lang_name = self.get_language_name(src_lang)
        return f"[Переведено с {lang_name}]\n\n{text}"

translator = TranslationManager()

# ========== РАСШИРЕННЫЙ ПОИСК В ИНТЕРНЕТЕ ==========
class WebSearchManager:
    """Менеджер для поиска в интернете с поддержкой 7 поисковиков"""
    
    def __init__(self):
        self.user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        ]
        self.session = None
        
    async def get_session(self):
        if not self.session:
            self.session = aiohttp.ClientSession()
        return self.session
    
    async def close(self):
        if self.session:
            await self.session.close()
    
    def _get_headers(self):
        return {
            'User-Agent': random.choice(self.user_agents),
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7',
            'DNT': '1',
            'Connection': 'keep-alive'
        }
    
    async def search_google(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://www.google.com/search?q={search_query}&num={num_results}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for g in soup.find_all('div', class_='g'):
                        title_elem = g.find('h3')
                        link_elem = g.find('a')
                        desc_elem = g.find('div', class_=['IsZvec', 'VwiC3b'])
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('/url?q='):
                                link = link.split('/url?q=')[1].split('&')[0]
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Google',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Google: {e}")
            return []
    
    async def search_bing(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://www.bing.com/search?q={search_query}&count={num_results}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for li in soup.find_all('li', class_='b_algo'):
                        title_elem = li.find('h2')
                        link_elem = li.find('a')
                        desc_elem = li.find('p')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Bing',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Bing: {e}")
            return []
    
    async def search_yahoo(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://search.yahoo.com/search?p={search_query}&n={num_results}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for div in soup.find_all('div', class_='algo'):
                        title_elem = div.find('h3')
                        link_elem = div.find('a')
                        desc_elem = div.find('div', class_='compText')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Yahoo',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Yahoo: {e}")
            return []
    
    async def search_yandex(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://yandex.ru/search/?text={search_query}&numdoc={num_results}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for li in soup.find_all('li', class_='serp-item'):
                        title_elem = li.find('h2')
                        link_elem = li.find('a')
                        desc_elem = li.find('div', class_='text-container')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Yandex',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Yandex: {e}")
            return []
    
    async def search_duckduckgo(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://html.duckduckgo.com/html/?q={search_query}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for result in soup.find_all('div', class_='result'):
                        title_elem = result.find('h2', class_='result__title')
                        link_elem = result.find('a', class_='result__a')
                        desc_elem = result.find('a', class_='result__snippet')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'DuckDuckGo',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка DuckDuckGo: {e}")
            return []
    
    async def search_baidu(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://www.baidu.com/s?wd={search_query}&rn={num_results}"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for div in soup.find_all('div', class_='result'):
                        title_elem = div.find('h3')
                        link_elem = div.find('a')
                        desc_elem = div.find('div', class_='c-abstract')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Baidu',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Baidu: {e}")
            return []
    
    async def search_ask(self, query: str, num_results: int = 20) -> List[dict]:
        try:
            results = []
            search_query = quote_plus(query)
            url = f"https://www.ask.com/web?q={search_query}&qo=pagination&page=1"
            
            session = await self.get_session()
            async with session.get(url, headers=self._get_headers()) as response:
                if response.status == 200:
                    html_content = await response.text()
                    soup = BeautifulSoup(html_content, 'html.parser')
                    
                    for div in soup.find_all('div', class_='PartialSearchResults-item'):
                        title_elem = div.find('a', class_='PartialSearchResults-item-title')
                        link_elem = title_elem
                        desc_elem = div.find('p', class_='PartialSearchResults-item-abstract')
                        
                        if title_elem and link_elem:
                            title = title_elem.get_text()
                            link = link_elem.get('href')
                            description = desc_elem.get_text() if desc_elem else ''
                            
                            if link and link.startswith('http'):
                                results.append({
                                    'title': title,
                                    'link': link,
                                    'description': description,
                                    'source': 'Ask.com',
                                    'language': translator.detect_language(title + ' ' + description)
                                })
                    
                    return results[:num_results]
        except Exception as e:
            print(f"Ошибка Ask.com: {e}")
            return []
    
    async def search_all(self, query: str, engines: List[str] = None, num_per_engine: int = 15) -> List[dict]:
        """Поиск во всех указанных поисковиках"""
        if engines is None:
            engines = ['google', 'bing', 'yahoo', 'yandex', 'duckduckgo', 'baidu', 'ask']
        
        all_results = []
        seen_links = set()
        
        search_tasks = []
        for engine in engines:
            if engine == 'google':
                search_tasks.append(self.search_google(query, num_per_engine))
            elif engine == 'bing':
                search_tasks.append(self.search_bing(query, num_per_engine))
            elif engine == 'yahoo':
                search_tasks.append(self.search_yahoo(query, num_per_engine))
            elif engine == 'yandex':
                search_tasks.append(self.search_yandex(query, num_per_engine))
            elif engine == 'duckduckgo':
                search_tasks.append(self.search_duckduckgo(query, num_per_engine))
            elif engine == 'baidu':
                search_tasks.append(self.search_baidu(query, num_per_engine))
            elif engine == 'ask':
                search_tasks.append(self.search_ask(query, num_per_engine))
        
        results_lists = await asyncio.gather(*search_tasks, return_exceptions=True)
        
        for results in results_lists:
            if isinstance(results, list):
                for result in results:
                    if result['link'] not in seen_links:
                        seen_links.add(result['link'])
                        all_results.append(result)
        
        return all_results[:MAX_WEB_RESULTS]

web_search = WebSearchManager()

# ========== ФУНКЦИИ ДЛЯ РАБОТЫ С ФАЙЛАМИ ==========
def load_channels():
    if os.path.exists(CHANNELS_FILE):
        with open(CHANNELS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_channels(channels):
    with open(CHANNELS_FILE, 'w', encoding='utf-8') as f:
        json.dump(channels, f, ensure_ascii=False, indent=2)

def normalize_channel_url(url):
    url = url.replace('@', '')
    
    if 't.me/' in url:
        parts = url.split('t.me/')
        username = parts[-1].strip('/')
        username = username.split('/')[0]
        return f"https://t.me/{username}"
    
    if '/' not in url and not url.startswith('http'):
        return f"https://t.me/{url}"
    
    return url

def extract_channel_name(url):
    url = normalize_channel_url(url)
    match = re.search(r't\.me/(?:s/)?([a-zA-Z0-9_]+)', url)
    if match:
        return match.group(1)
    return None

def export_channels_to_excel():
    try:
        channels = load_channels()
        df = pd.DataFrame(channels)
        df.insert(0, '№', range(1, len(df) + 1))
        filename = f"channels_export_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        filepath = os.path.join(UPLOADS_DIR, filename)
        df.to_excel(filepath, index=False, engine='openpyxl')
        return filepath
    except Exception as e:
        print(f"Ошибка экспорта: {e}")
        return None

def import_channels_from_excel(file_path):
    try:
        df = pd.read_excel(file_path, engine='openpyxl')
        current_channels = load_channels()
        current_urls = [ch['url'] for ch in current_channels]
        new_channels = []
        duplicates = []
        invalid = []
        
        for index, row in df.iterrows():
            name_col = None
            url_col = None
            
            for col in df.columns:
                col_lower = str(col).lower()
                if 'название' in col_lower or 'name' in col_lower:
                    name_col = col
                if 'ссылка' in col_lower or 'url' in col_lower:
                    url_col = col
            
            if name_col and url_col:
                name = str(row[name_col]).strip()
                url = str(row[url_col]).strip()
                
                if url and ('t.me/' in url or '@' in url):
                    url = normalize_channel_url(url)
                    if url in current_urls:
                        duplicates.append(f"{name} - {url}")
                    else:
                        new_channels.append({'name': name, 'url': url})
                        current_urls.append(url)
                else:
                    invalid.append(f"{name} - {url}")
        
        if new_channels:
            current_channels.extend(new_channels)
            save_channels(current_channels)
        
        return {
            'success': True,
            'added': len(new_channels),
            'duplicates': duplicates,
            'invalid': invalid,
            'total': len(current_channels)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def import_channels_from_txt(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        current_channels = load_channels()
        current_urls = [ch['url'] for ch in current_channels]
        new_channels = []
        duplicates = []
        invalid = []
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('#'):
                continue
            
            parts = line.split(',')
            if len(parts) >= 2:
                name = parts[0].strip()
                url = parts[1].strip()
            else:
                url = line
                name = extract_channel_name(url) or url.split('/')[-1]
            
            if url and ('t.me/' in url or '@' in url):
                url = normalize_channel_url(url)
                if url in current_urls:
                    duplicates.append(f"{name} - {url}")
                else:
                    new_channels.append({'name': name, 'url': url})
                    current_urls.append(url)
            else:
                invalid.append(f"{name} - {url}")
        
        if new_channels:
            current_channels.extend(new_channels)
            save_channels(current_channels)
        
        return {
            'success': True,
            'added': len(new_channels),
            'duplicates': duplicates,
            'invalid': invalid,
            'total': len(current_channels)
        }
    except Exception as e:
        return {'success': False, 'error': str(e)}

def safe_remove_file(filename):
    for i in range(5):
        try:
            if os.path.exists(filename):
                os.remove(filename)
                return True
        except PermissionError:
            time.sleep(0.5)
        except Exception:
            pass
    return False

def save_session_string(session_string):
    global MASTER_SESSION
    try:
        session_file = os.path.join(SESSIONS_DIR, f'master_session.txt')
        with open(session_file, 'w', encoding='utf-8') as f:
            f.write(session_string)
        MASTER_SESSION = session_string
        return True
    except Exception as e:
        print(f"Ошибка сохранения сессии: {e}")
        return False

def load_master_session():
    global MASTER_SESSION
    try:
        session_file = os.path.join(SESSIONS_DIR, f'master_session.txt')
        if os.path.exists(session_file):
            with open(session_file, 'r', encoding='utf-8') as f:
                MASTER_SESSION = f.read().strip()
                return MASTER_SESSION
    except Exception as e:
        print(f"Ошибка загрузки сессии: {e}")
    return None

def remove_master_session():
    global MASTER_SESSION
    try:
        session_file = os.path.join(SESSIONS_DIR, f'master_session.txt')
        if os.path.exists(session_file):
            os.remove(session_file)
        MASTER_SESSION = None
        return True
    except Exception:
        pass
    return False

def cleanup_temp_files(user_id):
    try:
        for file in os.listdir(IMAGES_DIR):
            if file.startswith(f"img_{user_id}_"):
                os.remove(os.path.join(IMAGES_DIR, file))
    except Exception as e:
        print(f"Ошибка очистки: {e}")

async def download_media(message, user_id, client):
    try:
        if message.media:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
            file_extension = '.jpg'
            
            if hasattr(message.media, 'document') and message.media.document:
                doc = message.media.document
                if doc.attributes:
                    for attr in doc.attributes:
                        if hasattr(attr, 'file_name') and attr.file_name:
                            file_extension = os.path.splitext(attr.file_name)[1]
                            break
                if not file_extension:
                    mime_to_ext = {
                        'image/jpeg': '.jpg', 'image/png': '.png', 'image/gif': '.gif',
                        'video/mp4': '.mp4', 'video/quicktime': '.mov',
                        'application/pdf': '.pdf', 'application/msword': '.doc',
                        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
                    }
                    file_extension = mime_to_ext.get(doc.mime_type, '.bin')
            
            local_filename = f"img_{user_id}_{timestamp}{file_extension}"
            local_path = os.path.join(IMAGES_DIR, local_filename)
            
            path = await message.download_media(file=local_path)
            return path if os.path.exists(path) else None
    except Exception as e:
        print(f"Ошибка скачивания медиа: {e}")
    return None

def add_image_to_doc(doc, image_path, max_width_inches=6):
    try:
        with Image.open(image_path) as img:
            width, height = img.size
            aspect = height / width
            doc_width = min(max_width_inches, width / 96)
            doc_height = doc_width * aspect
            doc.add_picture(image_path, width=Cm(doc_width * 2.54))
            return True
    except Exception as e:
        print(f"Ошибка добавления изображения: {e}")
        return False

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
    
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)
    r.append(rPr)
    
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    
    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink

def is_admin(user_id):
    return user_id == ADMIN_ID

def parse_period(text):
    text = text.lower().strip().rstrip('.')
    
    patterns = [
        (r'^(\d+)\s*(?:час|часа|часов|ч)$', 1),
        (r'^(\d+)\s*(?:минут|минута|минуты|мин|м)$', 1/60),
        (r'^(\d+)\s*(?:день|дня|дней|д|сут|суток)$', 24),
    ]
    
    for pattern, multiplier in patterns:
        match = re.match(pattern, text)
        if match:
            value = int(match.group(1))
            return value * multiplier
    
    match = re.match(r'^(\d+)$', text)
    if match:
        return int(match.group(1))
    
    return None

def format_period(period_hours):
    if period_hours < 1:
        minutes = int(period_hours * 60)
        if minutes == 1:
            return "1 минута"
        elif minutes < 5:
            return f"{minutes} минуты"
        else:
            return f"{minutes} минут"
    elif period_hours < 24:
        if period_hours == 1:
            return "1 час"
        elif period_hours < 5:
            return f"{int(period_hours)} часа"
        else:
            return f"{int(period_hours)} часов"
    else:
        days = period_hours / 24
        if days == 1:
            return "1 день"
        elif days < 5:
            return f"{int(days)} дня"
        else:
            return f"{int(days)} дней"

def format_datetime_utc10(dt):
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    dt_utc10 = dt.astimezone(UTC_PLUS_10)
    return dt_utc10.strftime('%d.%m.%Y %H:%M')

# ========== КЛАВИАТУРЫ ==========
def get_main_keyboard(user_id):
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("📋 Список каналов"),
        KeyboardButton("➕ Добавить канал"),
        KeyboardButton("📤 Импорт каналов"),
        KeyboardButton("📥 Экспорт каналов"),
        KeyboardButton("🔍 Расширенный поиск"),
        KeyboardButton("❓ Помощь"),
        KeyboardButton("⏹️ Стоп"),
        KeyboardButton("🛡️ Безопасность")
    )
    
    if is_admin(user_id):
        keyboard.add(
            KeyboardButton("🔄 Собрать всё"),
            KeyboardButton("🚪 Разлогиниться")
        )
    
    return keyboard

def get_search_type_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("📱 По каналам (из базы)"),
        KeyboardButton("🌍 По всему Telegram"),
        KeyboardButton("🔎 Поиск в интернете"),
        KeyboardButton("⚡ Комбинированный поиск"),
        KeyboardButton("◀️ Назад в меню")
    )
    return keyboard

def get_web_search_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("🔍 Все поисковики (7)"),
        KeyboardButton("🔍 Google"),
        KeyboardButton("🔎 Bing"),
        KeyboardButton("🔍 Yahoo"),
        KeyboardButton("🌍 Yandex"),
        KeyboardButton("🦆 DuckDuckGo"),
        KeyboardButton("🐼 Baidu"),
        KeyboardButton("❓ Ask.com"),
        KeyboardButton("◀️ Назад")
    )
    return keyboard

def get_period_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("🕐 1 час"),
        KeyboardButton("🕒 3 часа"),
        KeyboardButton("🕖 7 часов"),
        KeyboardButton("📅 24 часа"),
        KeyboardButton("📆 3 дня"),
        KeyboardButton("📆 7 дней"),
        KeyboardButton("⌨️ Свой период"),
        KeyboardButton("◀️ Назад")
    )
    return keyboard

def get_custom_period_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=1)
    keyboard.add(KeyboardButton("◀️ Назад"))
    return keyboard

def get_image_option_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("🖼️ С картинками"),
        KeyboardButton("📝 Только текст"),
        KeyboardButton("◀️ Назад")
    )
    return keyboard

def get_auth_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=1)
    keyboard.add(
        KeyboardButton("❌ Отменить авторизацию"),
        KeyboardButton("◀️ Назад в меню")
    )
    return keyboard

def get_channels_management_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("📤 Экспорт в Excel"),
        KeyboardButton("📥 Импорт из Excel"),
        KeyboardButton("📥 Импорт из TXT"),
        KeyboardButton("◀️ Назад в меню")
    )
    return keyboard

def get_translation_keyboard():
    keyboard = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    keyboard.add(
        KeyboardButton("🇷🇺 На русский"),
        KeyboardButton("🇬🇧 На английский"),
        KeyboardButton("🇩🇪 На немецкий"),
        KeyboardButton("🇫🇷 На французский"),
        KeyboardButton("🇪🇸 На испанский"),
        KeyboardButton("🇨🇳 На китайский"),
        KeyboardButton("🔄 Без перевода"),
        KeyboardButton("◀️ Назад")
    )
    return keyboard

# ========== КОМАНДЫ ==========
@dp.message_handler(commands=['start'])
async def cmd_start(message: types.Message):
    if not security.validate_telegram_request(message):
        return
    
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов. Подождите минуту.")
        return
    
    channels = load_channels()
    
    welcome_text = (
        "👋 Добро пожаловать!\n\n"
        "Я бот для сбора информации из Telegram-каналов и интернета.\n\n"
        "🔍 *Новые возможности:*\n"
        "• Поиск в 7 поисковых системах (Google, Bing, Yahoo, Yandex, DuckDuckGo, Baidu, Ask)\n"
        "• Автоматический перевод на разные языки с пометкой\n"
        "• Мощная система безопасности\n\n"
        "Как пользоваться:\n"
        "1️⃣ Нажми '➕ Добавить канал' для добавления канала\n"
        "2️⃣ Нажми '🔍 Расширенный поиск' и выбери тип поиска\n"
        "3️⃣ Введи ключевые слова\n"
        "4️⃣ Выбери язык перевода (если нужно)\n"
        "5️⃣ Выбери период времени\n"
        "6️⃣ Выбери формат отчета\n\n"
        "⏹️ Остановить - '⏹️ Стоп'\n"
        "🛡️ Безопасность - отчет о защите"
    )
    
    welcome_text += f"\n\n📌 Всего каналов в базе: {len(channels)}"
    
    if is_admin(user_id):
        welcome_text += (
            "\n\n👑 Вы администратор. Доступны доп. функции:\n"
            "• 🔄 Собрать всё\n"
            "• 🚪 Разлогиниться\n"
            "• 🛡️ Безопасность - подробный отчет"
        )
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        MASTER_SESSION = load_master_session()
    
    if not MASTER_SESSION and is_admin(user_id):
        welcome_text += "\n\n⚠️ Мастер-сессия не найдена! Нажми '🔍 Расширенный поиск' для авторизации."
    
    await message.reply(welcome_text, reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🛡️ Безопасность")
async def show_security(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if not is_admin(user_id):
        await message.reply("❌ Только для администратора.")
        return
    
    report = security.get_security_report()
    
    text = (
        "🛡️ *Отчет по безопасности*\n\n"
        f"📊 Всего логов: {report['total_logs']}\n"
        f"📈 Логов за час: {report['logs_last_hour']}\n"
        f"🚫 Заблокировано пользователей: {report['blocked_users']}\n"
        f"🌐 Заблокировано IP: {report['blocked_ips']}\n"
        f"⚠️ Нарушений rate limit: {report['rate_limit_violations']}\n"
        f"🔒 Брутфорс-атак: {report['brute_force_attempts']}\n\n"
        f"⚙️ Настройки:\n"
        f"• Rate limiting: {'✅' if SECURITY_SETTINGS['rate_limiting'] else '❌'}\n"
        f"• Защита от брутфорса: {'✅' if SECURITY_SETTINGS['brute_force_protection'] else '❌'}\n"
        f"• Защита от SQL: {'✅' if SECURITY_SETTINGS['sql_injection_protection'] else '❌'}\n"
        f"• Защита от XSS: {'✅' if SECURITY_SETTINGS['xss_protection'] else '❌'}\n"
    )
    
    await message.reply(text, parse_mode='Markdown')

@dp.message_handler(commands=['reset'])
async def cmd_reset(message: types.Message):
    user_id = message.from_user.id
    
    if not is_admin(user_id):
        await message.reply("❌ Только для администратора.")
        return
    
    if user_id in stop_flags:
        stop_flags[user_id] = True
    
    if user_id in auth_data and 'client' in auth_data[user_id]:
        try:
            await auth_data[user_id]['client'].disconnect()
        except:
            pass
    
    remove_master_session()
    cleanup_temp_files(user_id)
    
    if user_id in auth_data:
        del auth_data[user_id]
    
    await message.reply(
        "🗑 Мастер-сессия сброшена!",
        reply_markup=get_main_keyboard(user_id)
    )

@dp.message_handler(commands=['debug'])
async def cmd_debug(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if is_admin(user_id):
        session_file = os.path.join(SESSIONS_DIR, f'master_session.txt')
        session_exists = os.path.exists(session_file)
        in_auth = user_id in auth_data
        is_stopping = user_id in stop_flags and stop_flags[user_id]
        channels = load_channels()
        
        debug_text = (
            f"🔍 Отладка (админ)\n\n"
            f"📁 Мастер-сессия: {'✅' if session_exists else '❌'}\n"
            f"📊 Каналов: {len(channels)}\n"
            f"🔄 В авторизации: {'✅' if in_auth else '❌'}\n"
            f"⏹️ Процесс остановлен: {'✅' if is_stopping else '❌'}\n"
            f"🆔 User ID: {user_id}"
        )
    else:
        is_stopping = user_id in stop_flags and stop_flags[user_id]
        channels = load_channels()
        debug_text = (
            f"🔍 Отладка\n\n"
            f"📊 Каналов: {len(channels)}\n"
            f"⏹️ Процесс остановлен: {'✅' if is_stopping else '❌'}\n"
            f"🆔 User ID: {user_id}"
        )
    
    await message.reply(debug_text)

@dp.message_handler(commands=['checksession'])
async def cmd_checksession(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        MASTER_SESSION = load_master_session()
    
    if not MASTER_SESSION:
        await message.reply("❌ Мастер-сессия не найдена.")
        return
    
    if is_admin(user_id):
        await message.reply("🔍 Проверяю мастер-сессию...")
        
        try:
            client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
            await client.connect()
            
            if await client.is_user_authorized():
                me = await client.get_me()
                await message.reply(
                    f"✅ Мастер-сессия работает!\n\n"
                    f"Имя: {me.first_name}\n"
                    f"Username: @{me.username}"
                )
                new_session_string = client.session.save()
                save_session_string(new_session_string)
            else:
                await message.reply("❌ Мастер-сессия не активна.")
                remove_master_session()
            
            await client.disconnect()
        except Exception as e:
            await message.reply(f"❌ Ошибка: {str(e)}")
            remove_master_session()
    else:
        await message.reply("✅ Мастер-сессия активна.")

@dp.message_handler(lambda message: message.text == "📋 Список каналов")
async def show_channels(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    channels = load_channels()
    
    if not channels:
        await message.reply(
            "📭 Пока нет добавленных каналов",
            reply_markup=get_main_keyboard(user_id)
        )
        return
    
    keyboard = InlineKeyboardMarkup(row_width=1)
    for i, channel in enumerate(channels):
        channel_button = InlineKeyboardButton(f"📢 {channel['name']}", url=channel['url'])
        if is_admin(user_id):
            delete_button = InlineKeyboardButton(f"❌ Удалить", callback_data=f"delete_{i}")
            keyboard.row(channel_button, delete_button)
        else:
            keyboard.add(channel_button)
    
    if is_admin(user_id):
        keyboard.add(InlineKeyboardButton("❌ Удалить все", callback_data="delete_all"))
    
    keyboard.add(InlineKeyboardButton("◀️ Назад", callback_data="back_to_main"))
    
    text = f"📋 Общий список каналов:\n\n"
    for i, channel in enumerate(channels, 1):
        text += f"{i}. {channel['name']}\n"
    text += f"\nВсего каналов: {len(channels)}"
    
    await message.reply(text, reply_markup=keyboard)

@dp.message_handler(lambda message: message.text == "➕ Добавить канал")
async def add_channel_prompt(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    await message.reply(
        "🔗 Отправь ссылку на Telegram-канал\n\n"
        "Примеры:\n"
        "• https://t.me/durov\n"
        "• @durov\n"
        "• durov"
    )
    
    user_data[user_id] = {'state': 'waiting_channel_link'}

@dp.message_handler(lambda message: message.text == "📤 Импорт каналов")
async def import_channels_prompt(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    await message.reply(
        "📤 Выберите способ импорта:",
        reply_markup=get_channels_management_keyboard()
    )

@dp.message_handler(lambda message: message.text == "📥 Экспорт каналов")
async def export_channels(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    status_msg = await message.reply("🔄 Создаю Excel файл...")
    
    try:
        filepath = export_channels_to_excel()
        
        if filepath and os.path.exists(filepath):
            with open(filepath, 'rb') as f:
                await bot.send_document(user_id, f, caption="📊 Список каналов")
            os.remove(filepath)
            await status_msg.delete()
        else:
            await status_msg.edit_text("❌ Ошибка")
    except Exception as e:
        await status_msg.edit_text(f"❌ Ошибка: {str(e)}")

@dp.message_handler(lambda message: message.text == "📥 Импорт из Excel")
async def import_excel_prompt(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    await message.reply(
        "📥 Отправьте Excel файл (.xlsx)",
        reply_markup=get_channels_management_keyboard()
    )
    user_data[user_id] = {'state': 'waiting_excel_file'}

@dp.message_handler(lambda message: message.text == "📥 Импорт из TXT")
async def import_txt_prompt(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    await message.reply(
        "📥 Отправьте текстовый файл (.txt)",
        reply_markup=get_channels_management_keyboard()
    )
    user_data[user_id] = {'state': 'waiting_txt_file'}

@dp.message_handler(content_types=ContentType.DOCUMENT)
async def handle_document(message: types.Message):
    user_id = message.from_user.id
    
    if not security.validate_telegram_request(message):
        return
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if user_id not in user_data:
        await message.reply("Сначала выберите действие", reply_markup=get_main_keyboard(user_id))
        return
    
    state = user_data[user_id].get('state')
    
    if state not in ['waiting_excel_file', 'waiting_txt_file']:
        return
    
    document = message.document
    file_name = document.file_name
    file_ext = os.path.splitext(file_name)[1].lower()
    
    if state == 'waiting_excel_file' and file_ext not in ['.xlsx', '.xls']:
        await message.reply("❌ Отправьте Excel файл", reply_markup=get_channels_management_keyboard())
        return
    
    if state == 'waiting_txt_file' and file_ext != '.txt':
        await message.reply("❌ Отправьте текстовый файл", reply_markup=get_channels_management_keyboard())
        return
    
    status_msg = await message.reply("🔄 Обрабатываю...")
    
    try:
        file_path = os.path.join(UPLOADS_DIR, f"{user_id}_{file_name}")
        await document.download(destination_file=file_path)
        
        if state == 'waiting_excel_file':
            result = import_channels_from_excel(file_path)
        else:
            result = import_channels_from_txt(file_path)
        
        safe_remove_file(file_path)
        
        if result['success']:
            response = f"✅ Импорт завершен!\n\n"
            response += f"📊 Добавлено: {result['added']}\n"
            response += f"📈 Всего: {result['total']}\n"
            
            if result['duplicates']:
                response += f"\n⚠️ Дубликаты: {len(result['duplicates'])}"
            if result['invalid']:
                response += f"\n❌ Некорректные: {len(result['invalid'])}"
        else:
            response = f"❌ Ошибка: {result['error']}"
        
        await status_msg.delete()
        await message.reply(response, reply_markup=get_main_keyboard(user_id))
        del user_data[user_id]
        
    except Exception as e:
        await status_msg.delete()
        await message.reply(f"❌ Ошибка: {str(e)}")
        safe_remove_file(file_path)

@dp.message_handler(lambda message: message.text == "🔍 Расширенный поиск")
async def extended_search(message: types.Message):
    user_id = message.from_user.id
    
    if not security.validate_telegram_request(message):
        return
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        MASTER_SESSION = load_master_session()
    
    if not MASTER_SESSION:
        if is_admin(user_id):
            await message.reply(
                "🔄 Мастер-сессия не найдена. Начинаю авторизацию...",
                reply_markup=get_auth_keyboard()
            )
            await start_authorization(user_id, message)
        else:
            await message.reply(
                "❌ Мастер-сессия не найдена. Доступен только поиск в интернете.",
                reply_markup=get_search_type_keyboard()
            )
            user_data[user_id] = {'state': 'waiting_search_type'}
        return
    
    await message.reply(
        "🔍 Выберите источник поиска:",
        reply_markup=get_search_type_keyboard()
    )
    user_data[user_id] = {'state': 'waiting_search_type'}

@dp.message_handler(lambda message: message.text == "📱 По каналам (из базы)")
async def search_channels_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    channels = load_channels()
    
    if not channels:
        await message.reply("❌ Сначала добавьте каналы", reply_markup=get_main_keyboard(user_id))
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'channels',
        'channels': channels
    }
    
    await message.reply("🔍 Введите ключевые слова:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🌍 По всему Telegram")
async def search_global_telegram(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        if is_admin(user_id):
            await message.reply(
                "🔄 Мастер-сессия не найдена. Начинаю авторизацию...",
                reply_markup=get_auth_keyboard()
            )
            await start_authorization(user_id, message)
        else:
            await message.reply("❌ Нужна мастер-сессия.", reply_markup=get_main_keyboard(user_id))
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'global_telegram'
    }
    
    await message.reply("🔍 Введите ключевые слова для глобального поиска:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔎 Поиск в интернете")
async def search_web(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    await message.reply(
        "🌐 Выберите поисковую систему (доступно 7):",
        reply_markup=get_web_search_keyboard()
    )
    user_data[user_id] = {'state': 'waiting_web_search_type'}

@dp.message_handler(lambda message: message.text == "⚡ Комбинированный поиск")
async def search_combined(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        if is_admin(user_id):
            await message.reply(
                "🔄 Мастер-сессия не найдена. Начинаю авторизацию...",
                reply_markup=get_auth_keyboard()
            )
            await start_authorization(user_id, message)
        else:
            await message.reply(
                "❌ Будет выполнен только поиск в интернете.",
                reply_markup=get_main_keyboard(user_id)
            )
            user_data[user_id] = {
                'state': 'waiting_keywords',
                'search_type': 'web_only'
            }
            await message.reply("🔍 Введите ключевые слова:")
        return
    
    channels = load_channels()
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'combined',
        'channels': channels
    }
    
    await message.reply("🔍 Введите ключевые слова для комбинированного поиска:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔍 Все поисковики (7)")
async def search_all_engines(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'all_engines'
    }
    
    await message.reply("🌐 Введите ключевые слова для поиска во всех 7 поисковиках:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔍 Google")
async def search_google_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'google'
    }
    await message.reply("🔍 Введите ключевые слова для Google:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔎 Bing")
async def search_bing_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'bing'
    }
    await message.reply("🔎 Введите ключевые слова для Bing:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔍 Yahoo")
async def search_yahoo_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'yahoo'
    }
    await message.reply("🔍 Введите ключевые слова для Yahoo:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🌍 Yandex")
async def search_yandex_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'yandex'
    }
    await message.reply("🌍 Введите ключевые слова для Yandex:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🦆 DuckDuckGo")
async def search_duckduckgo_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'duckduckgo'
    }
    await message.reply("🦆 Введите ключевые слова для DuckDuckGo:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🐼 Baidu")
async def search_baidu_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'baidu'
    }
    await message.reply("🐼 Введите ключевые слова для Baidu:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "❓ Ask.com")
async def search_ask_only(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    user_data[user_id] = {
        'state': 'waiting_keywords',
        'search_type': 'ask'
    }
    await message.reply("❓ Введите ключевые слова для Ask.com:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🔄 Собрать всё")
async def collect_all(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if not is_admin(user_id):
        await message.reply("❌ Только для администратора.", reply_markup=get_main_keyboard(user_id))
        return
    
    global MASTER_SESSION
    if not MASTER_SESSION:
        MASTER_SESSION = load_master_session()
    
    if not MASTER_SESSION:
        await message.reply(
            "🔄 Мастер-сессия не найдена. Начинаю авторизацию...",
            reply_markup=get_auth_keyboard()
        )
        await start_authorization(user_id, message)
        return
    
    channels = load_channels()
    
    if not channels:
        await message.reply("❌ Сначала добавьте каналы", reply_markup=get_main_keyboard(user_id))
        return
    
    await message.reply(
        "⏱ Выбери период времени",
        reply_markup=get_period_keyboard()
    )
    
    user_data[user_id] = {
        'state': 'waiting_period',
        'keywords': 'все',
        'search_type': 'channels',
        'channels': channels
    }

@dp.message_handler(lambda message: message.text == "❓ Помощь")
async def show_help(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    channels = load_channels()
    
    help_text = (
        "❓ Помощь\n\n"
        "📋 Список каналов - посмотреть каналы\n"
        "➕ Добавить канал - добавить канал\n"
        "📤 Импорт/📥 Экспорт - управление каналами\n"
        "🔍 Расширенный поиск - выбор источника\n"
        "⏹️ Стоп - остановить поиск\n"
        "🛡️ Безопасность - отчет о защите\n\n"
        f"📊 Каналов в базе: {len(channels)}\n\n"
        "🔍 Типы поиска:\n"
        "• 📱 По каналам из базы\n"
        "• 🌍 По всему Telegram\n"
        "• 🔎 Поиск в интернете (7 поисковиков)\n"
        "• ⚡ Комбинированный\n\n"
        "🌐 Поисковики: Google, Bing, Yahoo, Yandex, DuckDuckGo, Baidu, Ask\n"
        "🌍 Перевод: автоматический с пометкой\n\n"
        "⌨️ Свой период: 30 минут, 2 часа, 5 дней и т.д.\n"
        f"Максимум: {MAX_PERIOD_DAYS} дней"
    )
    
    if is_admin(user_id):
        help_text += (
            "\n\n👑 Админ:\n"
            "🔄 Собрать всё\n"
            "🚪 Разлогиниться\n"
            "/reset - сброс"
        )
    
    await message.reply(help_text, reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "⏹️ Стоп")
async def stop_process(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if user_id in stop_flags:
        stop_flags[user_id] = True
        await message.reply("⏹️ Останавливаю...", reply_markup=get_main_keyboard(user_id))
    else:
        await message.reply("Нет активного процесса.", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "❌ Отменить авторизацию")
async def cancel_auth(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if not is_admin(user_id):
        return
    
    if user_id in auth_data:
        if 'client' in auth_data[user_id]:
            try:
                await auth_data[user_id]['client'].disconnect()
            except:
                pass
        del auth_data[user_id]
    
    await message.reply("❌ Авторизация отменена.", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "🚪 Разлогиниться")
async def logout(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if not is_admin(user_id):
        await message.reply("❌ Только для администратора.", reply_markup=get_main_keyboard(user_id))
        return
    
    if user_id in stop_flags:
        stop_flags[user_id] = True
    
    if user_id in auth_data and 'client' in auth_data[user_id]:
        try:
            await auth_data[user_id]['client'].disconnect()
        except:
            pass
    
    remove_master_session()
    
    if user_id in auth_data:
        del auth_data[user_id]
    
    await message.reply("🚪 Вы вышли.", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "◀️ Назад")
async def back_to_previous(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if user_id in user_data:
        current_state = user_data[user_id].get('state')
        
        if current_state == 'waiting_search_type':
            del user_data[user_id]
            await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))
        elif current_state == 'waiting_web_search_type':
            user_data[user_id]['state'] = 'waiting_search_type'
            await message.reply("🔍 Выберите источник поиска:", reply_markup=get_search_type_keyboard())
        elif current_state in ['waiting_keywords', 'waiting_period']:
            del user_data[user_id]
            await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))
        else:
            del user_data[user_id]
            await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))
    else:
        await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.text == "◀️ Назад в меню")
async def back_to_menu(message: types.Message):
    user_id = message.from_user.id
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if user_id in user_data:
        del user_data[user_id]
    
    await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_period')
async def process_period(message: types.Message):
    user_id = message.from_user.id
    period_text = message.text
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    period_hours = 0
    
    if period_text == "🕐 1 час":
        period_hours = 1
    elif period_text == "🕒 3 часа":
        period_hours = 3
    elif period_text == "🕖 7 часов":
        period_hours = 7
    elif period_text == "📅 24 часа":
        period_hours = 24
    elif period_text == "📆 3 дня":
        period_hours = 72
    elif period_text == "📆 7 дней":
        period_hours = 168
    elif period_text == "⌨️ Свой период":
        await message.reply(
            f"⌨️ Введите свой период\n\n"
            f"Примеры: 30 минут, 2 часа, 5 дней",
            reply_markup=get_custom_period_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_custom_period'
        return
    elif period_text == "◀️ Назад":
        if 'search_type' in user_data[user_id]:
            await message.reply("🔍 Выберите источник:", reply_markup=get_search_type_keyboard())
            user_data[user_id]['state'] = 'waiting_search_type'
        else:
            await message.reply("Главное меню:", reply_markup=get_main_keyboard(user_id))
            del user_data[user_id]
        return
    else:
        await message.reply("❌ Выберите из кнопок", reply_markup=get_period_keyboard())
        return
    
    user_data[user_id]['period_hours'] = period_hours
    user_data[user_id]['period_text'] = period_text
    
    # Если перевод доступен, предлагаем выбрать язык
    if TRANSLATION_AVAILABLE:
        await message.reply(
            "🌍 Выберите язык перевода (или без перевода):",
            reply_markup=get_translation_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_translation'
    else:
        user_data[user_id]['translation_lang'] = 'ru'
        await message.reply(
            "🖼️ Выберите формат отчета:",
            reply_markup=get_image_option_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_image_option'

@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_custom_period')
async def process_custom_period(message: types.Message):
    user_id = message.from_user.id
    period_text = message.text
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if period_text == "◀️ Назад":
        await message.reply("⏱ Выберите период:", reply_markup=get_period_keyboard())
        user_data[user_id]['state'] = 'waiting_period'
        return
    
    period_hours = parse_period(period_text)
    
    if period_hours is None or period_hours <= 0:
        await message.reply(
            "❌ Не удалось распознать период. Попробуйте еще раз:",
            reply_markup=get_custom_period_keyboard()
        )
        return
    
    if period_hours > MAX_PERIOD_HOURS:
        await message.reply(
            f"❌ Максимум {MAX_PERIOD_DAYS} дней. Введите меньший период:",
            reply_markup=get_custom_period_keyboard()
        )
        return
    
    user_data[user_id]['period_hours'] = period_hours
    user_data[user_id]['period_text'] = period_text
    
    if TRANSLATION_AVAILABLE:
        await message.reply(
            "🌍 Выберите язык перевода:",
            reply_markup=get_translation_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_translation'
    else:
        user_data[user_id]['translation_lang'] = 'ru'
        await message.reply(
            "🖼️ Выберите формат отчета:",
            reply_markup=get_image_option_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_image_option'

@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_translation')
async def process_translation(message: types.Message):
    user_id = message.from_user.id
    option = message.text
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if option == "◀️ Назад":
        await message.reply("⏱ Выберите период:", reply_markup=get_period_keyboard())
        user_data[user_id]['state'] = 'waiting_period'
        return
    
    translation_map = {
        "🇷🇺 На русский": 'ru',
        "🇬🇧 На английский": 'en',
        "🇩🇪 На немецкий": 'de',
        "🇫🇷 На французский": 'fr',
        "🇪🇸 На испанский": 'es',
        "🇨🇳 На китайский": 'zh-cn',
        "🔄 Без перевода": None
    }
    
    if option in translation_map:
        user_data[user_id]['translation_lang'] = translation_map[option]
        await message.reply(
            "🖼️ Выберите формат отчета:",
            reply_markup=get_image_option_keyboard()
        )
        user_data[user_id]['state'] = 'waiting_image_option'
    else:
        await message.reply("❌ Выберите из кнопок", reply_markup=get_translation_keyboard())

@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_image_option')
async def process_image_option(message: types.Message):
    user_id = message.from_user.id
    option = message.text
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if option == "◀️ Назад":
        if TRANSLATION_AVAILABLE:
            await message.reply("🌍 Выберите язык перевода:", reply_markup=get_translation_keyboard())
            user_data[user_id]['state'] = 'waiting_translation'
        else:
            await message.reply("⏱ Выберите период:", reply_markup=get_period_keyboard())
            user_data[user_id]['state'] = 'waiting_period'
        return
    elif option not in ["🖼️ С картинками", "📝 Только текст"]:
        await message.reply("❌ Выберите из кнопок", reply_markup=get_image_option_keyboard())
        return
    
    save_images = (option == "🖼️ С картинками")
    user_data[user_id]['save_images'] = save_images
    
    search_type = user_data[user_id].get('search_type', 'channels')
    
    if search_type == 'channels':
        await message.reply(
            f"🔍 Начинаю поиск по каналам\n"
            f"Период: {user_data[user_id]['period_text']}\n"
            f"Формат: {'с картинками' if save_images else 'только текст'}\n\n"
            f"⏳ Это займет время...",
            reply_markup=get_main_keyboard(user_id)
        )
        await collect_from_channels(user_id)
    
    elif search_type == 'global_telegram':
        await message.reply(
            f"🌍 Начинаю глобальный поиск по Telegram\n"
            f"Период: {user_data[user_id]['period_text']}\n"
            f"Формат: {'с картинками' if save_images else 'только текст'}",
            reply_markup=get_main_keyboard(user_id)
        )
        await collect_global_telegram(user_id)
    
    elif search_type in ['google', 'bing', 'yahoo', 'yandex', 'duckduckgo', 'baidu', 'ask', 'all_engines']:
        await message.reply(
            f"🔎 Начинаю поиск в интернете\n"
            f"Поисковик: {search_type}\n"
            f"Формат: {'с картинками' if save_images else 'только текст'}",
            reply_markup=get_main_keyboard(user_id)
        )
        await collect_from_web(user_id)
    
    elif search_type == 'combined':
        await message.reply(
            f"⚡ Начинаю комбинированный поиск\n"
            f"• По каналам из базы\n"
            f"• Глобальный поиск по Telegram\n"
            f"• Поиск в интернете (7 поисковиков)\n"
            f"Период: {user_data[user_id]['period_text']}\n"
            f"Формат: {'с картинками' if save_images else 'только текст'}",
            reply_markup=get_main_keyboard(user_id)
        )
        await collect_combined(user_id)
    
    elif search_type == 'web_only':
        await message.reply(
            f"🔎 Начинаю поиск в интернете\n"
            f"Формат: {'с картинками' if save_images else 'только текст'}",
            reply_markup=get_main_keyboard(user_id)
        )
        await collect_from_web(user_id)

# ========== АВТОРИЗАЦИЯ ==========
async def start_authorization(user_id, message):
    if not is_admin(user_id):
        return
    
    if not security.check_brute_force(user_id, 'auth'):
        await message.reply("❌ Слишком много попыток. Подождите.")
        return
    
    try:
        if user_id in auth_data and 'client' in auth_data[user_id]:
            try:
                await auth_data[user_id]['client'].disconnect()
            except:
                pass
        
        client = TelegramClient(StringSession(), API_ID, API_HASH)
        await client.connect()
        
        auth_data[user_id] = {
            'client': client,
            'state': 'waiting_phone'
        }
        
        await message.reply(
            "📱 Введи номер телефона (например: +79123456789):",
            reply_markup=get_auth_keyboard()
        )
        
    except Exception as e:
        await message.reply(f"❌ Ошибка: {str(e)}")
        if user_id in auth_data:
            del auth_data[user_id]

async def start_authorization_without_message(user_id):
    if not is_admin(user_id):
        return
    
    try:
        if user_id in auth_data and 'client' in auth_data[user_id]:
            try:
                await auth_data[user_id]['client'].disconnect()
            except:
                pass
        
        client = TelegramClient(StringSession(), API_ID, API_HASH)
        await client.connect()
        
        auth_data[user_id] = {
            'client': client,
            'state': 'waiting_phone'
        }
        
        await bot.send_message(
            user_id, 
            "📱 Введи номер телефона:",
            reply_markup=get_auth_keyboard()
        )
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {str(e)}")
        if user_id in auth_data:
            del auth_data[user_id]

@dp.message_handler(lambda message: message.from_user.id in auth_data and auth_data[message.from_user.id].get('state') == 'waiting_phone')
async def process_phone(message: types.Message):
    user_id = message.from_user.id
    
    if not is_admin(user_id):
        return
    
    if not security.check_brute_force(user_id, 'auth'):
        await message.reply("❌ Слишком много попыток.")
        return
    
    phone = message.text.strip()
    
    if phone == "❌ Отменить авторизацию":
        await cancel_auth(message)
        return
    
    # Валидация номера
    valid, error = security.validate_input(phone, 20)
    if not valid:
        await message.reply(f"❌ {error}")
        return
    
    phone = re.sub(r'[^\d+]', '', phone)
    
    status_msg = await message.reply("📲 Отправляю код...")
    
    try:
        client = auth_data[user_id]['client']
        result = await client.send_code_request(phone)
        
        auth_data[user_id]['phone'] = phone
        auth_data[user_id]['phone_code_hash'] = result.phone_code_hash
        auth_data[user_id]['state'] = 'waiting_code'
        
        await status_msg.delete()
        await message.reply(
            "🔐 Введи код из Telegram (только цифры):",
            reply_markup=get_auth_keyboard()
        )
        
    except FloodWaitError as e:
        await status_msg.delete()
        await message.reply(f"⏳ Подождите {e.seconds} секунд")
    except Exception as e:
        await status_msg.delete()
        await message.reply(f"❌ Ошибка: {str(e)[:200]}")
        await client.disconnect()
        if user_id in auth_data:
            del auth_data[user_id]

@dp.message_handler(lambda message: message.from_user.id in auth_data and auth_data[message.from_user.id].get('state') == 'waiting_code')
async def process_code(message: types.Message):
    user_id = message.from_user.id
    
    if not is_admin(user_id):
        return
    
    if not security.check_brute_force(user_id, 'code'):
        await message.reply("❌ Слишком много попыток ввода кода.")
        return
    
    code = message.text.strip()
    
    if code == "❌ Отменить авторизацию":
        await cancel_auth(message)
        return
    
    # Валидация кода
    valid, error = security.validate_input(code, 10)
    if not valid:
        await message.reply(f"❌ {error}")
        return
    
    code = re.sub(r'\D', '', code)
    
    status_msg = await message.reply("🔐 Проверяю код...")
    
    try:
        client = auth_data[user_id]['client']
        phone = auth_data[user_id]['phone']
        
        await client.sign_in(phone, code)
        
        await status_msg.delete()
        
        session_string = client.session.save()
        save_session_string(session_string)
        
        me = await client.get_me()
        
        await message.reply(
            f"✅ *Авторизация успешна!*\n\n"
            f"Аккаунт: {me.first_name}\n"
            f"Username: @{me.username}",
            parse_mode='Markdown',
            reply_markup=get_main_keyboard(user_id)
        )
        
        if user_id in auth_data:
            del auth_data[user_id]
        
    except SessionPasswordNeededError:
        await status_msg.delete()
        auth_data[user_id]['state'] = 'waiting_password'
        await message.reply(
            "🔐 Требуется пароль двухфакторки. Введи пароль:",
            reply_markup=get_auth_keyboard()
        )
    except PhoneCodeExpiredError:
        await status_msg.delete()
        await message.reply("❌ Код истек. Введи номер заново:")
        auth_data[user_id]['state'] = 'waiting_phone'
    except PhoneCodeInvalidError:
        await status_msg.delete()
        await message.reply("❌ Неверный код. Попробуй еще раз:")
    except Exception as e:
        await status_msg.delete()
        await message.reply(f"❌ Ошибка: {str(e)[:200]}")
        await client.disconnect()
        if user_id in auth_data:
            del auth_data[user_id]

@dp.message_handler(lambda message: message.from_user.id in auth_data and auth_data[message.from_user.id].get('state') == 'waiting_password')
async def process_password(message: types.Message):
    user_id = message.from_user.id
    
    if not is_admin(user_id):
        return
    
    password = message.text.strip()
    
    if password == "❌ Отменить авторизацию":
        await cancel_auth(message)
        return
    
    try:
        client = auth_data[user_id]['client']
        await client.sign_in(password=password)
        
        session_string = client.session.save()
        save_session_string(session_string)
        
        me = await client.get_me()
        
        await message.reply(
            f"✅ *Авторизация успешна!*\n\n"
            f"Аккаунт: {me.first_name}",
            parse_mode='Markdown',
            reply_markup=get_main_keyboard(user_id)
        )
        
        if user_id in auth_data:
            del auth_data[user_id]
        
    except Exception as e:
        await message.reply(f"❌ Неверный пароль. Попробуй еще раз:", reply_markup=get_auth_keyboard())

# ========== ОБРАБОТЧИКИ ТЕКСТА ==========
@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_channel_link')
async def process_channel_link(message: types.Message):
    try:
        user_id = message.from_user.id
        link = message.text.strip()
        
        # Валидация ссылки
        valid, error = security.validate_input(link, 200)
        if not valid:
            await message.reply(f"❌ {error}")
            return
        
        link = normalize_channel_url(link)
        channel_name = extract_channel_name(link) or link.split('/')[-1]
        
        status_msg = await message.reply("🔄 Проверяю канал...")
        
        try:
            global MASTER_SESSION
            if MASTER_SESSION:
                client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
                await client.connect()
                
                if await client.is_user_authorized():
                    entity = await client.get_entity(link)
                    if hasattr(entity, 'title') and entity.title:
                        channel_name = entity.title
                
                await client.disconnect()
        except:
            pass
        
        await status_msg.delete()
        
        channels = load_channels()
        
        for ch in channels:
            if ch['url'] == link or ch['name'].lower() == channel_name.lower():
                await message.reply(
                    f"❌ Канал уже есть!",
                    reply_markup=get_main_keyboard(user_id)
                )
                del user_data[user_id]
                return
        
        channels.append({'name': channel_name, 'url': link})
        save_channels(channels)
        
        del user_data[user_id]
        
        await message.reply(
            f"✅ Канал добавлен!\n\n"
            f"Название: {channel_name}\n"
            f"Ссылка: {link}\n"
            f"Всего каналов: {len(channels)}",
            reply_markup=get_main_keyboard(user_id)
        )
        
    except Exception as e:
        await message.reply(f"❌ Ошибка: {str(e)}")

@dp.message_handler(lambda message: message.from_user.id in user_data and user_data[message.from_user.id].get('state') == 'waiting_keywords')
async def process_keywords(message: types.Message):
    keywords = message.text.strip()
    user_id = message.from_user.id
    
    # Валидация ключевых слов
    valid, error = security.validate_input(keywords, 200)
    if not valid:
        await message.reply(f"❌ {error}")
        return
    
    user_data[user_id]['keywords'] = keywords
    
    await message.reply(
        "⏱ Выбери период времени:",
        reply_markup=get_period_keyboard()
    )
    user_data[user_id]['state'] = 'waiting_period'

# ========== СБОР ДАННЫХ ==========
async def collect_from_channels(user_id):
    client = None
    try:
        stop_flags[user_id] = False
        
        keywords = user_data[user_id]['keywords']
        channels = user_data[user_id].get('channels', load_channels())
        period_hours = user_data[user_id].get('period_hours', 168)
        save_images = user_data[user_id].get('save_images', True)
        translation_lang = user_data[user_id].get('translation_lang', 'ru')
        
        await bot.send_message(user_id, "🔄 Подключаюсь к Telegram...")
        
        global MASTER_SESSION
        if not MASTER_SESSION:
            MASTER_SESSION = load_master_session()
        
        if not MASTER_SESSION:
            await bot.send_message(user_id, "❌ Мастер-сессия не найдена.", reply_markup=get_main_keyboard(user_id))
            return
        
        client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
        await client.connect()
        
        if not await client.is_user_authorized():
            await bot.send_message(user_id, "❌ Мастер-сессия не активна.", reply_markup=get_main_keyboard(user_id))
            await client.disconnect()
            remove_master_session()
            return
        
        await bot.send_message(user_id, "✅ Подключено! Начинаю сбор...")
        
        doc = Document()
        
        title = doc.add_heading('Отчёт по Telegram-каналам', 0)
        title.alignment = 1
        
        now_utc10 = datetime.now(UTC_PLUS_10)
        doc.add_paragraph(f"Дата: {now_utc10.strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Тип: По каналам из базы")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
        if translation_lang:
            doc.add_paragraph(f"Перевод: {'Да' if translation_lang else 'Нет'}")
        doc.add_paragraph()
        
        total_posts = 0
        processed_channels = 0
        stopped_early = False
        channels_with_posts = 0
        
        now = datetime.now().astimezone()
        start_time = now - timedelta(hours=period_hours)
        
        for channel in channels:
            if stop_flags.get(user_id, False):
                stopped_early = True
                await bot.send_message(user_id, f"⏹️ Остановлено. Обработано: {processed_channels}/{len(channels)}")
                break
            
            processed_channels += 1
            await bot.send_message(user_id, f"📱 [{processed_channels}/{len(channels)}] {channel['name']}")
            
            try:
                channel_url = normalize_channel_url(channel['url'])
                entity = await client.get_entity(channel_url)
                
                posts_count = 0
                channel_posts = []
                
                async for message in client.iter_messages(entity, offset_date=now, reverse=False):
                    if posts_count % 5 == 0 and stop_flags.get(user_id, False):
                        stopped_early = True
                        break
                    
                    if message.date:
                        msg_date = message.date
                        if msg_date.tzinfo is None:
                            msg_date = msg_date.replace(tzinfo=timezone.utc)
                        if msg_date < start_time:
                            break
                    
                    if keywords.lower() == 'все' or (message.text and keywords.lower() in message.text.lower()):
                        channel_posts.append(message)
                        posts_count += 1
                        total_posts += 1
                        
                        if posts_count % 20 == 0:
                            await bot.send_message(user_id, f"📊 Найдено {posts_count} постов...")
                
                if posts_count > 0:
                    channels_with_posts += 1
                    doc.add_heading(f"Канал: {channel['name']}", level=1)
                    doc.add_paragraph(f"Ссылка: {channel_url}")
                    doc.add_paragraph()
                    
                    for message in channel_posts:
                        p = doc.add_paragraph()
                        
                        display_date = format_datetime_utc10(message.date)
                        p.add_run(f"📅 {display_date}\n").bold = True
                        
                        if message.text:
                            text = message.text
                            if len(text) > 5000:
                                text = text[:5000] + "..."
                            
                            # Перевод текста, если нужно
                            if translation_lang:
                                trans_result = await translator.translate(text, translation_lang)
                                if trans_result['was_translated']:
                                    p.add_run(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                                else:
                                    p.add_run(text)
                            else:
                                p.add_run(text)
                        
                        if save_images and message.media:
                            if hasattr(message.media, 'photo'):
                                image_path = await download_media(message, user_id, client)
                                if image_path:
                                    doc.add_paragraph()
                                    add_image_to_doc(doc, image_path)
                            
                            if hasattr(message.media, 'document') and message.media.document:
                                mime = message.media.document.mime_type
                                if mime and 'image' in mime:
                                    image_path = await download_media(message, user_id, client)
                                    if image_path:
                                        doc.add_paragraph()
                                        add_image_to_doc(doc, image_path)
                        
                        if message.id:
                            channel_username = channel_url.split('/')[-1]
                            link_url = f"https://t.me/{channel_username}/{message.id}"
                            link_paragraph = doc.add_paragraph()
                            add_hyperlink(link_paragraph, "🔗 Ссылка", link_url)
                        
                        doc.add_paragraph()
                    
                    doc.add_paragraph(f"✅ Найдено: {posts_count}")
                    doc.add_page_break()
                
                if stop_flags.get(user_id, False):
                    stopped_early = True
                    break
                
            except FloodWaitError as e:
                wait_time = e.seconds
                await bot.send_message(user_id, f"⚠️ Лимит. Жду {wait_time} сек...")
                await asyncio.sleep(wait_time)
                continue
            except Exception as e:
                await bot.send_message(user_id, f"⚠️ Ошибка: {str(e)[:100]}")
                doc.add_paragraph(f"❌ Ошибка: {channel['name']}")
                doc.add_page_break()
        
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f"Обработано каналов: {processed_channels}/{len(channels)}")
        doc.add_paragraph(f"Каналов с постами: {channels_with_posts}")
        doc.add_paragraph(f"Всего постов: {total_posts}")
        doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
        doc.add_paragraph(f"Формат: {'с картинками' if save_images else 'только текст'}")
        if stopped_early:
            doc.add_paragraph("⚠️ Остановлено досрочно")
        
        if total_posts == 0:
            await bot.send_message(
                user_id, 
                f"📭 Ничего не найдено по запросу: {keywords}",
                reply_markup=get_main_keyboard(user_id)
            )
            if user_id in user_data:
                del user_data[user_id]
            await client.disconnect()
            return
        
        output_file = f"report_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(output_file)
        
        with open(output_file, 'rb') as f:
            caption = f"✅ Отчет готов!\n\n"
            caption += f"📊 Постов: {total_posts}\n"
            caption += f"📁 Каналов: {channels_with_posts}/{processed_channels}\n"
            caption += f"🔍 Слова: {keywords}\n"
            caption += f"⏱ Период: {user_data[user_id].get('period_text', format_period(period_hours))}\n"
            if translation_lang:
                caption += f"🌍 Перевод: {translation_lang}\n"
            if stopped_early:
                caption += f"⚠️ Остановлено\n"
            
            await bot.send_document(user_id, f, caption=caption)
        
        if os.path.exists(output_file):
            os.remove(output_file)
        
        cleanup_temp_files(user_id)
        
        if user_id in user_data:
            del user_data[user_id]
        
        await client.disconnect()
        
        if user_id in stop_flags:
            del stop_flags[user_id]
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {str(e)}")
        logging.error(f"Error: {e}")
        cleanup_temp_files(user_id)
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]
        if client:
            await client.disconnect()

async def collect_global_telegram(user_id):
    """Глобальный поиск по всему Telegram (исправленная версия)"""
    client = None
    try:
        stop_flags[user_id] = False
        
        keywords = user_data[user_id]['keywords']
        period_hours = user_data[user_id].get('period_hours', 168)
        save_images = user_data[user_id].get('save_images', True)
        translation_lang = user_data[user_id].get('translation_lang', 'ru')
        
        await bot.send_message(user_id, "🔄 Подключаюсь для глобального поиска...")
        
        global MASTER_SESSION
        if not MASTER_SESSION:
            MASTER_SESSION = load_master_session()
        
        if not MASTER_SESSION:
            await bot.send_message(user_id, "❌ Мастер-сессия не найдена.", reply_markup=get_main_keyboard(user_id))
            return
        
        client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
        await client.connect()
        
        if not await client.is_user_authorized():
            await bot.send_message(user_id, "❌ Мастер-сессия не активна.", reply_markup=get_main_keyboard(user_id))
            await client.disconnect()
            remove_master_session()
            return
        
        await bot.send_message(user_id, "✅ Подключено! Выполняю поиск...")
        
        doc = Document()
        
        title = doc.add_heading('Глобальный поиск в Telegram', 0)
        title.alignment = 1
        
        now_utc10 = datetime.now(UTC_PLUS_10)
        doc.add_paragraph(f"Дата: {now_utc10.strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
        if translation_lang:
            doc.add_paragraph(f"Перевод: {'Да' if translation_lang else 'Нет'}")
        doc.add_paragraph()
        
        start_time = datetime.now().astimezone() - timedelta(hours=period_hours)
        
        await bot.send_message(user_id, "🌍 Ищу по всему Telegram...")
        
        try:
            result = await client(SearchGlobalRequest(
                q=keywords,
                filter=InputMessagesFilterEmpty(),
                min_date=start_time,
                max_date=datetime.now().astimezone(),
                offset_rate=0,
                offset_peer=None,
                offset_id=0,
                limit=100
            ))
            
            total_posts = 0
            
            if hasattr(result, 'messages') and result.messages:
                for msg in result.messages:
                    if stop_flags.get(user_id, False):
                        break
                    
                    # Пропускаем сообщения без текста
                    if not hasattr(msg, 'message') or not msg.message:
                        continue
                    
                    total_posts += 1
                    
                    # Получаем информацию о чате с проверкой на None
                    chat_title = 'Неизвестный чат'
                    chat_username = None
                    
                    try:
                        if msg.peer_id:  # Проверяем, что peer_id не None
                            chat = await client.get_entity(msg.peer_id)
                            chat_title = getattr(chat, 'title', 'Неизвестный чат')
                            chat_username = getattr(chat, 'username', None)
                        else:
                            # Если peer_id None, пропускаем это сообщение
                            continue
                    except Exception as e:
                        # Если не удалось получить информацию, используем заглушку
                        chat_title = 'Неизвестный чат'
                        chat_username = None
                    
                    doc.add_heading(f"Чат: {chat_title}", level=2)
                    if chat_username:
                        doc.add_paragraph(f"Ссылка: https://t.me/{chat_username}")
                    
                    p = doc.add_paragraph()
                    display_date = format_datetime_utc10(msg.date)
                    p.add_run(f"📅 {display_date}\n").bold = True
                    
                    if msg.message:
                        text = msg.message
                        if len(text) > 5000:
                            text = text[:5000] + "..."
                        
                        # Перевод текста
                        if translation_lang:
                            trans_result = await translator.translate(text, translation_lang)
                            if trans_result['was_translated']:
                                p.add_run(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                            else:
                                p.add_run(text)
                        else:
                            p.add_run(text)
                    
                    if save_images and msg.media:
                        if hasattr(msg.media, 'photo'):
                            image_path = await download_media(msg, user_id, client)
                            if image_path:
                                doc.add_paragraph()
                                add_image_to_doc(doc, image_path)
                    
                    if chat_username and msg.id:
                        link_url = f"https://t.me/{chat_username}/{msg.id}"
                        link_paragraph = doc.add_paragraph()
                        add_hyperlink(link_paragraph, "🔗 Ссылка", link_url)
                    
                    doc.add_paragraph()
                    
                    if total_posts % 10 == 0:
                        await bot.send_message(user_id, f"📊 Найдено {total_posts}...")
            
            doc.add_heading('Статистика', level=1)
            doc.add_paragraph(f"Всего сообщений: {total_posts}")
            doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
            doc.add_paragraph(f"Формат: {'с картинками' if save_images else 'только текст'}")
            
            if total_posts == 0:
                await bot.send_message(
                    user_id, 
                    f"📭 Ничего не найдено: {keywords}",
                    reply_markup=get_main_keyboard(user_id)
                )
                if user_id in user_data:
                    del user_data[user_id]
                await client.disconnect()
                return
            
            output_file = f"global_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
            doc.save(output_file)
            
            with open(output_file, 'rb') as f:
                caption = f"✅ Глобальный поиск завершен!\n\n"
                caption += f"📊 Найдено: {total_posts}\n"
                caption += f"🔍 Слова: {keywords}\n"
                caption += f"⏱ Период: {user_data[user_id].get('period_text', format_period(period_hours))}\n"
                if translation_lang:
                    caption += f"🌍 Перевод: {translation_lang}\n"
                
                await bot.send_document(user_id, f, caption=caption)
            
            if os.path.exists(output_file):
                os.remove(output_file)
            
        except Exception as e:
            await bot.send_message(user_id, f"❌ Ошибка поиска: {str(e)}")
            doc.add_paragraph(f"❌ Ошибка: {str(e)}")
        
        cleanup_temp_files(user_id)
        
        if user_id in user_data:
            del user_data[user_id]
        
        await client.disconnect()
        
        if user_id in stop_flags:
            del stop_flags[user_id]
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {str(e)}")
        logging.error(f"Error: {e}")
        cleanup_temp_files(user_id)
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]
        if client:
            await client.disconnect()

async def collect_from_web(user_id):
    """Сбор данных из интернета (7 поисковиков)"""
    try:
        stop_flags[user_id] = False
        
        keywords = user_data[user_id]['keywords']
        search_type = user_data[user_id].get('search_type', 'all_engines')
        save_images = user_data[user_id].get('save_images', True)
        translation_lang = user_data[user_id].get('translation_lang', 'ru')
        
        await bot.send_message(user_id, "🌐 Начинаю поиск в интернете...")
        
        doc = Document()
        
        title = doc.add_heading('Поиск в интернете', 0)
        title.alignment = 1
        
        now_utc10 = datetime.now(UTC_PLUS_10)
        doc.add_paragraph(f"Дата: {now_utc10.strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        if translation_lang:
            doc.add_paragraph(f"Перевод: {'Да' if translation_lang else 'Нет'}")
        doc.add_paragraph()
        
        total_results = 0
        
        # Определяем, какие поисковики использовать
        engines_to_use = []
        if search_type == 'all_engines':
            engines_to_use = ['google', 'bing', 'yahoo', 'yandex', 'duckduckgo', 'baidu', 'ask']
        else:
            engines_to_use = [search_type]
        
        # Поиск в каждом поисковике
        for engine in engines_to_use:
            if stop_flags.get(user_id, False):
                break
            
            await bot.send_message(user_id, f"🔍 Ищу в {engine.capitalize()}...")
            
            if engine == 'google':
                results = await web_search.search_google(keywords, 15)
            elif engine == 'bing':
                results = await web_search.search_bing(keywords, 15)
            elif engine == 'yahoo':
                results = await web_search.search_yahoo(keywords, 15)
            elif engine == 'yandex':
                results = await web_search.search_yandex(keywords, 15)
            elif engine == 'duckduckgo':
                results = await web_search.search_duckduckgo(keywords, 15)
            elif engine == 'baidu':
                results = await web_search.search_baidu(keywords, 15)
            elif engine == 'ask':
                results = await web_search.search_ask(keywords, 15)
            else:
                results = []
            
            if results:
                doc.add_heading(f'{engine.capitalize()}', level=1)
                for result in results:
                    if stop_flags.get(user_id, False):
                        break
                    
                    doc.add_heading(result['title'], level=2)
                    doc.add_paragraph(f"Источник: {result['source']}")
                    
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, result['link'], result['link'])
                    
                    if result['description']:
                        desc = result['description']
                        
                        # Перевод описания
                        if translation_lang:
                            trans_result = await translator.translate(desc, translation_lang)
                            if trans_result['was_translated']:
                                doc.add_paragraph(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                            else:
                                doc.add_paragraph(desc)
                        else:
                            doc.add_paragraph(desc)
                    
                    doc.add_paragraph()
                    total_results += 1
                
                doc.add_page_break()
        
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f"Всего результатов: {total_results}")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        doc.add_paragraph(f"Поисковики: {', '.join(engines_to_use)}")
        
        if total_results == 0:
            await bot.send_message(
                user_id, 
                f"📭 Ничего не найдено: {keywords}",
                reply_markup=get_main_keyboard(user_id)
            )
            if user_id in user_data:
                del user_data[user_id]
            return
        
        output_file = f"web_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(output_file)
        
        with open(output_file, 'rb') as f:
            caption = f"✅ Поиск в интернете завершен!\n\n"
            caption += f"📊 Результатов: {total_results}\n"
            caption += f"🔍 Слова: {keywords}\n"
            caption += f"🌐 Поисковики: {len(engines_to_use)}\n"
            if translation_lang:
                caption += f"🌍 Перевод: {translation_lang}\n"
            
            await bot.send_document(user_id, f, caption=caption)
        
        if os.path.exists(output_file):
            os.remove(output_file)
        
        cleanup_temp_files(user_id)
        
        if user_id in user_data:
            del user_data[user_id]
        
        if user_id in stop_flags:
            del stop_flags[user_id]
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {str(e)}")
        logging.error(f"Error: {e}")
        cleanup_temp_files(user_id)
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

async def collect_combined(user_id):
    """Комбинированный поиск"""
    try:
        stop_flags[user_id] = False
        
        keywords = user_data[user_id]['keywords']
        period_hours = user_data[user_id].get('period_hours', 168)
        save_images = user_data[user_id].get('save_images', True)
        translation_lang = user_data[user_id].get('translation_lang', 'ru')
        
        await bot.send_message(user_id, "⚡ Начинаю комбинированный поиск...")
        
        doc = Document()
        
        title = doc.add_heading('Комбинированный отчёт', 0)
        title.alignment = 1
        
        now_utc10 = datetime.now(UTC_PLUS_10)
        doc.add_paragraph(f"Дата: {now_utc10.strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
        if translation_lang:
            doc.add_paragraph(f"Перевод: {'Да' if translation_lang else 'Нет'}")
        doc.add_paragraph()
        
        total_results = 0
        
        # 1. Поиск в интернете (7 поисковиков)
        if not stop_flags.get(user_id, False):
            await bot.send_message(user_id, "🌐 Ищу в интернете...")
            doc.add_heading('ИНТЕРНЕТ', level=1)
            
            all_web_results = await web_search.search_all(keywords, num_per_engine=5)
            if all_web_results:
                for result in all_web_results[:30]:  # Ограничиваем до 30
                    if stop_flags.get(user_id, False):
                        break
                    
                    doc.add_heading(result['title'], level=2)
                    doc.add_paragraph(f"Источник: {result['source']}")
                    
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, result['link'], result['link'])
                    
                    if result['description']:
                        desc = result['description']
                        if translation_lang:
                            trans_result = await translator.translate(desc, translation_lang)
                            if trans_result['was_translated']:
                                doc.add_paragraph(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                            else:
                                doc.add_paragraph(desc)
                        else:
                            doc.add_paragraph(desc)
                    
                    doc.add_paragraph()
                    total_results += 1
                
                doc.add_page_break()
        
        # 2. Глобальный поиск по Telegram
        if not stop_flags.get(user_id, False) and MASTER_SESSION:
            await bot.send_message(user_id, "🌍 Ищу по всему Telegram...")
            doc.add_heading('TELEGRAM (ГЛОБАЛЬНЫЙ)', level=1)
            
            client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
            await client.connect()
            
            if await client.is_user_authorized():
                start_time = datetime.now().astimezone() - timedelta(hours=period_hours)
                
                try:
                    result = await client(SearchGlobalRequest(
                        q=keywords,
                        filter=InputMessagesFilterEmpty(),
                        min_date=start_time,
                        max_date=datetime.now().astimezone(),
                        offset_rate=0,
                        offset_peer=None,
                        offset_id=0,
                        limit=30
                    ))
                    
                    telegram_count = 0
                    if hasattr(result, 'messages') and result.messages:
                        for msg in result.messages:
                            if stop_flags.get(user_id, False):
                                break
                            
                            if not hasattr(msg, 'message') or not msg.message:
                                continue
                            
                            telegram_count += 1
                            
                            chat_title = 'Неизвестный чат'
                            chat_username = None
                            
                            try:
                                if msg.peer_id:
                                    chat = await client.get_entity(msg.peer_id)
                                    chat_title = getattr(chat, 'title', 'Неизвестный чат')
                                    chat_username = getattr(chat, 'username', None)
                                else:
                                    continue
                            except:
                                pass
                            
                            doc.add_heading(f"Чат: {chat_title}", level=2)
                            if chat_username:
                                doc.add_paragraph(f"https://t.me/{chat_username}")
                            
                            p = doc.add_paragraph()
                            display_date = format_datetime_utc10(msg.date)
                            p.add_run(f"📅 {display_date}\n").bold = True
                            
                            if msg.message:
                                text = msg.message[:1000]
                                if translation_lang:
                                    trans_result = await translator.translate(text, translation_lang)
                                    if trans_result['was_translated']:
                                        p.add_run(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                                    else:
                                        p.add_run(text)
                                else:
                                    p.add_run(text)
                            
                            if chat_username and msg.id:
                                link_url = f"https://t.me/{chat_username}/{msg.id}"
                                link_paragraph = doc.add_paragraph()
                                add_hyperlink(link_paragraph, "🔗 Ссылка", link_url)
                            
                            doc.add_paragraph()
                    
                    doc.add_paragraph(f"✅ Найдено: {telegram_count}")
                    total_results += telegram_count
                    
                except Exception as e:
                    doc.add_paragraph(f"❌ Ошибка: {str(e)}")
            
            await client.disconnect()
            doc.add_page_break()
        
        # 3. Поиск по каналам из базы
        if not stop_flags.get(user_id, False) and MASTER_SESSION:
            await bot.send_message(user_id, "📱 Ищу по каналам...")
            doc.add_heading('КАНАЛЫ ИЗ БАЗЫ', level=1)
            
            channels = load_channels()
            if channels:
                client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
                await client.connect()
                
                if await client.is_user_authorized():
                    start_time = datetime.now().astimezone() - timedelta(hours=period_hours)
                    channels_count = 0
                    
                    for channel in channels[:3]:  # Ограничиваем до 3 каналов
                        if stop_flags.get(user_id, False):
                            break
                        
                        try:
                            channel_url = normalize_channel_url(channel['url'])
                            entity = await client.get_entity(channel_url)
                            
                            channel_posts = 0
                            async for message in client.iter_messages(entity, limit=5, offset_date=datetime.now()):
                                if stop_flags.get(user_id, False):
                                    break
                                
                                if message.date and message.date >= start_time and message.text and keywords.lower() in message.text.lower():
                                    if channel_posts == 0:
                                        doc.add_heading(f"{channel['name']}", level=2)
                                        doc.add_paragraph(f"Ссылка: {channel_url}")
                                    
                                    channel_posts += 1
                                    
                                    p = doc.add_paragraph()
                                    display_date = format_datetime_utc10(message.date)
                                    p.add_run(f"📅 {display_date}\n").bold = True
                                    
                                    text = message.text[:1000]
                                    if translation_lang:
                                        trans_result = await translator.translate(text, translation_lang)
                                        if trans_result['was_translated']:
                                            p.add_run(translator.add_translation_note(trans_result['translated'], trans_result['src_lang']))
                                        else:
                                            p.add_run(text)
                                    else:
                                        p.add_run(text)
                                    
                                    if message.id:
                                        channel_username = channel_url.split('/')[-1]
                                        link_url = f"https://t.me/{channel_username}/{message.id}"
                                        link_paragraph = doc.add_paragraph()
                                        add_hyperlink(link_paragraph, "🔗 Ссылка", link_url)
                                    
                                    doc.add_paragraph()
                            
                            if channel_posts > 0:
                                doc.add_paragraph(f"✅ Найдено: {channel_posts}")
                                channels_count += 1
                                total_results += channel_posts
                            
                        except Exception as e:
                            continue
                    
                    doc.add_paragraph(f"📊 Каналов с результатами: {channels_count}")
                
                await client.disconnect()
        
        doc.add_heading('ИТОГО', level=1)
        doc.add_paragraph(f"Всего результатов: {total_results}")
        doc.add_paragraph(f"Ключевые слова: {keywords}")
        doc.add_paragraph(f"Период: {user_data[user_id].get('period_text', format_period(period_hours))}")
        
        if total_results == 0:
            await bot.send_message(
                user_id, 
                f"📭 Ничего не найдено: {keywords}",
                reply_markup=get_main_keyboard(user_id)
            )
            if user_id in user_data:
                del user_data[user_id]
            return
        
        output_file = f"combined_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        doc.save(output_file)
        
        with open(output_file, 'rb') as f:
            caption = f"✅ Комбинированный поиск завершен!\n\n"
            caption += f"📊 Всего: {total_results}\n"
            caption += f"🔍 Слова: {keywords}\n"
            caption += f"⏱ Период: {user_data[user_id].get('period_text', format_period(period_hours))}\n"
            if translation_lang:
                caption += f"🌍 Перевод: {translation_lang}\n"
            
            await bot.send_document(user_id, f, caption=caption)
        
        if os.path.exists(output_file):
            os.remove(output_file)
        
        cleanup_temp_files(user_id)
        
        if user_id in user_data:
            del user_data[user_id]
        
        if user_id in stop_flags:
            del stop_flags[user_id]
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {str(e)}")
        logging.error(f"Error: {e}")
        cleanup_temp_files(user_id)
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

# ========== КОЛБЭКИ ==========
@dp.callback_query_handler(lambda c: c.data.startswith('delete_'))
async def process_delete_callback(callback_query: types.CallbackQuery):
    await bot.answer_callback_query(callback_query.id)
    
    user_id = callback_query.from_user.id
    
    if not security.verify_owner(user_id):
        await bot.send_message(user_id, "❌ Только для администратора.")
        return
    
    data = callback_query.data
    channels = load_channels()
    
    if data == "delete_all":
        save_channels([])
        await bot.send_message(user_id, "🗑 Все каналы удалены", reply_markup=get_main_keyboard(user_id))
        return
    
    index = int(data.split('_')[1])
    if index < len(channels):
        deleted = channels.pop(index)
        save_channels(channels)
        await bot.send_message(
            user_id,
            f"🗑 Удален: {deleted['name']}\nОсталось: {len(channels)}",
            reply_markup=get_main_keyboard(user_id)
        )

@dp.callback_query_handler(lambda c: c.data == 'back_to_main')
async def process_back_callback(callback_query: types.CallbackQuery):
    await bot.answer_callback_query(callback_query.id)
    user_id = callback_query.from_user.id
    
    if user_id in user_data:
        del user_data[user_id]
    
    await bot.send_message(user_id, "Главное меню:", reply_markup=get_main_keyboard(user_id))

# ========== НЕИЗВЕСТНЫЕ КОМАНДЫ ==========
@dp.message_handler()
async def handle_unknown(message: types.Message):
    user_id = message.from_user.id
    
    if not security.validate_telegram_request(message):
        return
    
    if not security.check_rate_limit(user_id):
        await message.reply("⏳ Слишком много запросов.")
        return
    
    if is_admin(user_id) and user_id in auth_data:
        if auth_data[user_id].get('state') == 'waiting_phone':
            await process_phone(message)
            return
        elif auth_data[user_id].get('state') == 'waiting_code':
            await process_code(message)
            return
        elif auth_data[user_id].get('state') == 'waiting_password':
            await process_password(message)
            return
    
    if message.text and ('t.me/' in message.text or '@' in message.text):
        user_data[user_id] = {'state': 'waiting_channel_link'}
        await process_channel_link(message)
    else:
        await message.reply(
            "Используй кнопки 👇",
            reply_markup=get_main_keyboard(user_id)
        )

# ========== ЗАПУСК ==========
if __name__ == '__main__':
    MASTER_SESSION = load_master_session()
    channels = load_channels()
    
    print("=" * 50)
    print("🤖 Бот запускается...")
    print(f"👑 Админ ID: {ADMIN_ID}")
    print(f"📊 Каналов: {len(channels)}")
    print(f"🔍 Поисковиков: 7 (Google, Bing, Yahoo, Yandex, DuckDuckGo, Baidu, Ask)")
    print(f"🌍 Перевод: {'Доступен' if TRANSLATION_AVAILABLE else 'Недоступен'}")
    print(f"🛡️ Безопасность: Включена")
    print("=" * 50)
    
    if MASTER_SESSION:
        print("✅ Мастер-сессия загружена")
    else:
        print("⚠️ Мастер-сессия не найдена")
    
    executor.start_polling(dp, skip_updates=True)
