import asyncio
import logging
from datetime import datetime, timedelta, timezone
import os
import json
import re
import time
import sys
import signal
import random

from PIL import Image
import pandas as pd
import openpyxl
import aiohttp
from bs4 import BeautifulSoup
import requests
from urllib.parse import quote_plus

# Для перевода
from googletrans import Translator, LANGUAGES
import langdetect

from aiogram import Bot, Dispatcher, types
from aiogram.types import ParseMode, ContentType, InlineKeyboardMarkup, InlineKeyboardButton, ReplyKeyboardMarkup, KeyboardButton
from aiogram.utils import executor
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from telethon import TelegramClient
from telethon.errors import FloodWaitError, SessionPasswordNeededError, PhoneCodeExpiredError, PhoneCodeInvalidError
from telethon.sessions import StringSession
from telethon.tl.functions.messages import SearchGlobalRequest
from telethon.tl.types import InputMessagesFilterEmpty

# ========== НАСТРОЙКИ ==========
API_TOKEN = '8029293386:AAG-Hih0eVun77YYcY8zf6PyDjQERmQCx9w'
API_ID = 38892524
API_HASH = 'd71ef1a657ab20d2a47a52130626c939'
ADMIN_ID = 5224743551

CHANNELS_FILE = 'saved_channels.json'
SESSIONS_DIR = 'telegram_sessions'
IMAGES_DIR = 'temp_images'
UPLOADS_DIR = 'uploads'

UTC_PLUS_10 = timezone(timedelta(hours=10))
MAX_PERIOD_DAYS = 30
MAX_PERIOD_HOURS = MAX_PERIOD_DAYS * 24

# Поисковики
SEARCH_ENGINES = ['google', 'bing', 'yahoo', 'yandex', 'duckduckgo']

os.makedirs(SESSIONS_DIR, exist_ok=True)
os.makedirs(IMAGES_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)

# ========== ИНИЦИАЛИЗАЦИЯ ==========
bot = Bot(token=API_TOKEN)
dp = Dispatcher(bot)
logging.basicConfig(level=logging.INFO)

user_data = {}
auth_data = {}
stop_flags = {}
MASTER_SESSION = None

# ========== ПЕРЕВОД ==========
class TranslatorManager:
    def __init__(self):
        self.translator = Translator()
        
    def detect_language(self, text):
        try:
            return langdetect.detect(text)
        except:
            return 'unknown'
    
    def get_language_name(self, code):
        return LANGUAGES.get(code, code)
    
    async def translate(self, text, dest='ru'):
        try:
            if not text:
                return {'translated': '', 'was_translated': False}
            
            src = self.detect_language(text)
            
            if src == dest:
                return {'translated': text, 'was_translated': False}
            
            result = await self.translator.translate(text, dest=dest)
            return {
                'translated': result.text,
                'src_lang': src,
                'src_name': self.get_language_name(src),
                'was_translated': True
            }
        except:
            return {'translated': text, 'was_translated': False}

translator = TranslatorManager()

# ========== ПОИСК В ИНТЕРНЕТЕ ==========
class WebSearch:
    def __init__(self):
        self.session = None
        
    async def get_session(self):
        if not self.session:
            self.session = aiohttp.ClientSession()
        return self.session
    
    async def close(self):
        if self.session:
            await self.session.close()
    
    async def search_google(self, query, num=10):
        try:
            results = []
            url = f"https://www.google.com/search?q={quote_plus(query)}&num={num}"
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            
            session = await self.get_session()
            async with session.get(url, headers=headers) as resp:
                if resp.status == 200:
                    html = await resp.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    for g in soup.find_all('div', class_='g'):
                        title = g.find('h3')
                        link = g.find('a')
                        desc = g.find('div', class_='VwiC3b')
                        
                        if title and link:
                            href = link.get('href')
                            if href and href.startswith('/url?q='):
                                href = href.split('/url?q=')[1].split('&')[0]
                            
                            if href and href.startswith('http'):
                                results.append({
                                    'title': title.text,
                                    'link': href,
                                    'description': desc.text if desc else '',
                                    'source': 'Google'
                                })
            return results
        except:
            return []
    
    async def search_bing(self, query, num=10):
        try:
            results = []
            url = f"https://www.bing.com/search?q={quote_plus(query)}&count={num}"
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            
            session = await self.get_session()
            async with session.get(url, headers=headers) as resp:
                if resp.status == 200:
                    html = await resp.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    for li in soup.find_all('li', class_='b_algo'):
                        title = li.find('h2')
                        link = li.find('a')
                        desc = li.find('p')
                        
                        if title and link:
                            href = link.get('href')
                            if href and href.startswith('http'):
                                results.append({
                                    'title': title.text,
                                    'link': href,
                                    'description': desc.text if desc else '',
                                    'source': 'Bing'
                                })
            return results
        except:
            return []
    
    async def search_yandex(self, query, num=10):
        try:
            results = []
            url = f"https://yandex.ru/search/?text={quote_plus(query)}&numdoc={num}"
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            
            session = await self.get_session()
            async with session.get(url, headers=headers) as resp:
                if resp.status == 200:
                    html = await resp.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    for li in soup.find_all('li', class_='serp-item'):
                        title = li.find('h2')
                        link = li.find('a')
                        desc = li.find('div', class_='text-container')
                        
                        if title and link:
                            href = link.get('href')
                            if href and href.startswith('http'):
                                results.append({
                                    'title': title.text,
                                    'link': href,
                                    'description': desc.text if desc else '',
                                    'source': 'Yandex'
                                })
            return results
        except:
            return []
    
    async def search_duckduckgo(self, query, num=10):
        try:
            results = []
            url = f"https://html.duckduckgo.com/html/?q={quote_plus(query)}"
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            
            session = await self.get_session()
            async with session.get(url, headers=headers) as resp:
                if resp.status == 200:
                    html = await resp.text()
                    soup = BeautifulSoup(html, 'html.parser')
                    
                    for result in soup.find_all('div', class_='result'):
                        title = result.find('h2', class_='result__title')
                        link = result.find('a', class_='result__a')
                        desc = result.find('a', class_='result__snippet')
                        
                        if title and link:
                            href = link.get('href')
                            if href and href.startswith('http'):
                                results.append({
                                    'title': title.text,
                                    'link': href,
                                    'description': desc.text if desc else '',
                                    'source': 'DuckDuckGo'
                                })
            return results
        except:
            return []

web_search = WebSearch()

# ========== РАБОТА С ФАЙЛАМИ ==========
def load_channels():
    if os.path.exists(CHANNELS_FILE):
        with open(CHANNELS_FILE, 'r', encoding='utf-8') as f:
            return json.load(f)
    return []

def save_channels(channels):
    with open(CHANNELS_FILE, 'w', encoding='utf-8') as f:
        json.dump(channels, f, ensure_ascii=False, indent=2)

def normalize_channel_url(url):
    url = url.replace('@', '')
    if 't.me/' in url:
        username = url.split('t.me/')[-1].strip('/').split('/')[0]
        return f"https://t.me/{username}"
    if '/' not in url and not url.startswith('http'):
        return f"https://t.me/{url}"
    return url

def extract_channel_name(url):
    url = normalize_channel_url(url)
    match = re.search(r't\.me/([a-zA-Z0-9_]+)', url)
    return match.group(1) if match else url.split('/')[-1]

def save_session_string(session_string):
    global MASTER_SESSION
    try:
        with open(os.path.join(SESSIONS_DIR, 'master_session.txt'), 'w', encoding='utf-8') as f:
            f.write(session_string)
        MASTER_SESSION = session_string
        return True
    except:
        return False

def load_master_session():
    global MASTER_SESSION
    try:
        path = os.path.join(SESSIONS_DIR, 'master_session.txt')
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                MASTER_SESSION = f.read().strip()
                return MASTER_SESSION
    except:
        pass
    return None

def remove_master_session():
    global MASTER_SESSION
    try:
        path = os.path.join(SESSIONS_DIR, 'master_session.txt')
        if os.path.exists(path):
            os.remove(path)
        MASTER_SESSION = None
    except:
        pass

def cleanup_temp_files(user_id):
    try:
        for f in os.listdir(IMAGES_DIR):
            if f.startswith(f"img_{user_id}_"):
                os.remove(os.path.join(IMAGES_DIR, f))
    except:
        pass

async def download_media(message, user_id, client):
    try:
        if message.media:
            filename = f"img_{user_id}_{datetime.now().timestamp()}.jpg"
            path = os.path.join(IMAGES_DIR, filename)
            return await message.download_media(file=path)
    except:
        pass
    return None

def add_image_to_doc(doc, path):
    try:
        doc.add_picture(path, width=Cm(10))
        return True
    except:
        return False

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True))
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)

def is_admin(user_id):
    return user_id == ADMIN_ID

def parse_period(text):
    text = text.lower().strip()
    patterns = [
        (r'(\d+)\s*(час|часа|часов|ч)', 1),
        (r'(\d+)\s*(минут|минута|минуты|мин|м)', 1/60),
        (r'(\d+)\s*(день|дня|дней|д|сут|суток)', 24),
    ]
    for pattern, mult in patterns:
        m = re.search(pattern, text)
        if m:
            return int(m.group(1)) * mult
    m = re.search(r'(\d+)', text)
    return int(m.group(1)) if m else None

def format_period(hours):
    if hours < 1:
        return f"{int(hours*60)} мин"
    if hours < 24:
        return f"{int(hours)} ч"
    return f"{int(hours/24)} дн"

def format_datetime_utc10(dt):
    if dt.tzinfo is None:
        dt = dt.replace(tzinfo=timezone.utc)
    return dt.astimezone(UTC_PLUS_10).strftime('%d.%m.%Y %H:%M')

# ========== КЛАВИАТУРЫ ==========
def get_main_keyboard(user_id):
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("📋 Список каналов"),
        KeyboardButton("➕ Добавить канал"),
        KeyboardButton("🔍 Поиск"),
        KeyboardButton("❓ Помощь"),
        KeyboardButton("⏹️ Стоп")
    )
    if is_admin(user_id):
        kb.add(KeyboardButton("🔄 Собрать всё"), KeyboardButton("🚪 Выйти"))
    return kb

def get_search_type_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("📱 По каналам"),
        KeyboardButton("🌍 По всему Telegram"),
        KeyboardButton("🔎 В интернете"),
        KeyboardButton("⚡ Комбинированный"),
        KeyboardButton("◀️ Назад")
    )
    return kb

def get_web_search_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("🔍 Все поисковики"),
        KeyboardButton("🔍 Google"),
        KeyboardButton("🔎 Bing"),
        KeyboardButton("🌍 Yandex"),
        KeyboardButton("🦆 DuckDuckGo"),
        KeyboardButton("◀️ Назад")
    )
    return kb

def get_period_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("🕐 1 час"),
        KeyboardButton("🕒 3 часа"),
        KeyboardButton("🕖 7 часов"),
        KeyboardButton("📅 24 часа"),
        KeyboardButton("📆 3 дня"),
        KeyboardButton("📆 7 дней"),
        KeyboardButton("⌨️ Свой"),
        KeyboardButton("◀️ Назад")
    )
    return kb

def get_translation_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("🇷🇺 На русский"),
        KeyboardButton("🇬🇧 На английский"),
        KeyboardButton("🇩🇪 На немецкий"),
        KeyboardButton("🇫🇷 На французский"),
        KeyboardButton("🇪🇸 На испанский"),
        KeyboardButton("🔄 Без перевода"),
        KeyboardButton("◀️ Назад")
    )
    return kb

def get_image_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=2)
    kb.add(
        KeyboardButton("🖼️ С картинками"),
        KeyboardButton("📝 Только текст"),
        KeyboardButton("◀️ Назад")
    )
    return kb

def get_auth_keyboard():
    kb = ReplyKeyboardMarkup(resize_keyboard=True, row_width=1)
    kb.add(KeyboardButton("❌ Отмена"), KeyboardButton("◀️ Назад"))
    return kb

# ========== КОМАНДЫ ==========
@dp.message_handler(commands=['start'])
async def start(message: types.Message):
    user_id = message.from_user.id
    channels = load_channels()
    
    text = (
        "👋 Привет!\n\n"
        "Я бот для поиска в Telegram и интернете\n\n"
        "🔍 Что умею:\n"
        "• Искать по твоим каналам\n"
        "• Искать по всему Telegram\n"
        "• Искать в Google, Bing, Yandex, DuckDuckGo\n"
        "• Переводить текст\n\n"
        f"📊 Каналов в базе: {len(channels)}\n"
    )
    
    if not MASTER_SESSION and is_admin(user_id):
        text += "\n⚠️ Нажми 🔍 Поиск для авторизации"
    
    await message.reply(text, reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text == "📋 Список каналов")
async def show_channels(m: types.Message):
    user_id = m.from_user.id
    channels = load_channels()
    
    if not channels:
        await m.reply("📭 Нет каналов", reply_markup=get_main_keyboard(user_id))
        return
    
    kb = InlineKeyboardMarkup(row_width=1)
    for i, ch in enumerate(channels):
        btn = InlineKeyboardButton(f"📢 {ch['name']}", url=ch['url'])
        if is_admin(user_id):
            del_btn = InlineKeyboardButton("❌", callback_data=f"del_{i}")
            kb.row(btn, del_btn)
        else:
            kb.add(btn)
    
    if is_admin(user_id):
        kb.add(InlineKeyboardButton("❌ Удалить все", callback_data="del_all"))
    
    text = "📋 Каналы:\n" + "\n".join(f"{i+1}. {ch['name']}" for i, ch in enumerate(channels))
    await m.reply(text, reply_markup=kb)

@dp.message_handler(lambda m: m.text == "➕ Добавить канал")
async def add_channel(m: types.Message):
    user_id = m.from_user.id
    await m.reply("🔗 Отправь ссылку на канал\nПример: @durov или https://t.me/durov")
    user_data[user_id] = {'state': 'waiting_channel'}

@dp.message_handler(lambda m: m.text == "🔍 Поиск")
async def search_menu(m: types.Message):
    user_id = m.from_user.id
    
    if not MASTER_SESSION:
        MASTER_SESSION = load_master_session()
    
    if not MASTER_SESSION and is_admin(user_id):
        await m.reply("🔄 Нужна авторизация", reply_markup=get_auth_keyboard())
        await start_auth(user_id, m)
        return
    
    await m.reply("🔍 Где ищем?", reply_markup=get_search_type_keyboard())
    user_data[user_id] = {'state': 'waiting_search_type'}

@dp.message_handler(lambda m: m.text == "📱 По каналам")
async def search_channels(m: types.Message):
    user_id = m.from_user.id
    channels = load_channels()
    
    if not channels:
        await m.reply("❌ Сначала добавь каналы", reply_markup=get_main_keyboard(user_id))
        return
    
    user_data[user_id] = {'state': 'waiting_keywords', 'search_type': 'channels', 'channels': channels}
    await m.reply("🔍 Введи ключевые слова:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text == "🌍 По всему Telegram")
async def search_global(m: types.Message):
    user_id = m.from_user.id
    user_data[user_id] = {'state': 'waiting_keywords', 'search_type': 'global'}
    await m.reply("🔍 Введи ключевые слова для поиска по всему Telegram:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text == "🔎 В интернете")
async def search_web_menu(m: types.Message):
    user_id = m.from_user.id
    await m.reply("🌐 Выбери поисковик:", reply_markup=get_web_search_keyboard())
    user_data[user_id] = {'state': 'waiting_web_type'}

@dp.message_handler(lambda m: m.text == "⚡ Комбинированный")
async def search_combined(m: types.Message):
    user_id = m.from_user.id
    channels = load_channels()
    user_data[user_id] = {'state': 'waiting_keywords', 'search_type': 'combined', 'channels': channels}
    await m.reply("🔍 Введи ключевые слова (поиск везде):", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text in ["🔍 Все поисковики", "🔍 Google", "🔎 Bing", "🌍 Yandex", "🦆 DuckDuckGo"])
async def web_type_selected(m: types.Message):
    user_id = m.from_user.id
    types = {
        "🔍 Все поисковики": "all",
        "🔍 Google": "google",
        "🔎 Bing": "bing",
        "🌍 Yandex": "yandex",
        "🦆 DuckDuckGo": "duckduckgo"
    }
    user_data[user_id] = {'state': 'waiting_keywords', 'search_type': 'web', 'web_type': types[m.text]}
    await m.reply("🔍 Введи ключевые слова:", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text == "🔄 Собрать всё")
async def collect_all(m: types.Message):
    user_id = m.from_user.id
    if not is_admin(user_id):
        await m.reply("❌ Только для админа")
        return
    
    channels = load_channels()
    if not channels:
        await m.reply("❌ Нет каналов")
        return
    
    user_data[user_id] = {'state': 'waiting_period', 'keywords': 'все', 'search_type': 'channels', 'channels': channels}
    await m.reply("⏱ За какой период?", reply_markup=get_period_keyboard())

@dp.message_handler(lambda m: m.text == "❓ Помощь")
async def help_cmd(m: types.Message):
    text = (
        "❓ Помощь\n\n"
        "📋 Список каналов - посмотреть каналы\n"
        "➕ Добавить канал - добавить канал\n"
        "🔍 Поиск - начать поиск\n"
        "⏹️ Стоп - остановить поиск\n\n"
        "🌐 Поисковики: Google, Bing, Yandex, DuckDuckGo\n"
        "🌍 Перевод: русский, английский, немецкий, французский, испанский\n"
        "⌨️ Свой период: 30 минут, 2 часа, 5 дней"
    )
    await m.reply(text)

@dp.message_handler(lambda m: m.text == "⏹️ Стоп")
async def stop_cmd(m: types.Message):
    user_id = m.from_user.id
    if user_id in stop_flags:
        stop_flags[user_id] = True
        await m.reply("⏹️ Останавливаю...")

@dp.message_handler(lambda m: m.text == "🚪 Выйти")
async def logout(m: types.Message):
    user_id = m.from_user.id
    if is_admin(user_id):
        remove_master_session()
        if user_id in auth_data:
            del auth_data[user_id]
        await m.reply("🚪 Вышел")

@dp.message_handler(lambda m: m.text == "❌ Отмена")
async def cancel(m: types.Message):
    user_id = m.from_user.id
    if user_id in auth_data:
        if 'client' in auth_data[user_id]:
            await auth_data[user_id]['client'].disconnect()
        del auth_data[user_id]
    await m.reply("❌ Отменено", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.text == "◀️ Назад")
async def back(m: types.Message):
    user_id = m.from_user.id
    
    if user_id in user_data:
        state = user_data[user_id].get('state')
        
        if state == 'waiting_search_type':
            del user_data[user_id]
            await m.reply("Главное меню", reply_markup=get_main_keyboard(user_id))
        elif state == 'waiting_web_type':
            user_data[user_id]['state'] = 'waiting_search_type'
            await m.reply("🔍 Где ищем?", reply_markup=get_search_type_keyboard())
        elif state in ['waiting_keywords', 'waiting_period']:
            del user_data[user_id]
            await m.reply("Главное меню", reply_markup=get_main_keyboard(user_id))
        else:
            del user_data[user_id]
            await m.reply("Главное меню", reply_markup=get_main_keyboard(user_id))
    else:
        await m.reply("Главное меню", reply_markup=get_main_keyboard(user_id))

# ========== АВТОРИЗАЦИЯ ==========
async def start_auth(user_id, message):
    try:
        if user_id in auth_data and 'client' in auth_data[user_id]:
            await auth_data[user_id]['client'].disconnect()
        
        client = TelegramClient(StringSession(), API_ID, API_HASH)
        await client.connect()
        
        auth_data[user_id] = {'client': client, 'state': 'waiting_phone'}
        await message.reply("📱 Введи номер телефона (например: +79123456789):", reply_markup=get_auth_keyboard())
    except Exception as e:
        await message.reply(f"❌ Ошибка: {e}")

@dp.message_handler(lambda m: m.from_user.id in auth_data and auth_data[m.from_user.id]['state'] == 'waiting_phone')
async def process_phone(m: types.Message):
    user_id = m.from_user.id
    phone = m.text.strip()
    
    if phone == "❌ Отмена":
        await cancel(m)
        return
    
    phone = re.sub(r'[^\d+]', '', phone)
    
    try:
        client = auth_data[user_id]['client']
        result = await client.send_code_request(phone)
        
        auth_data[user_id]['phone'] = phone
        auth_data[user_id]['hash'] = result.phone_code_hash
        auth_data[user_id]['state'] = 'waiting_code'
        
        await m.reply("🔐 Введи код из Telegram:", reply_markup=get_auth_keyboard())
    except Exception as e:
        await m.reply(f"❌ Ошибка: {e}")
        await client.disconnect()
        del auth_data[user_id]

@dp.message_handler(lambda m: m.from_user.id in auth_data and auth_data[m.from_user.id]['state'] == 'waiting_code')
async def process_code(m: types.Message):
    user_id = m.from_user.id
    code = m.text.strip()
    
    if code == "❌ Отмена":
        await cancel(m)
        return
    
    code = re.sub(r'\D', '', code)
    
    try:
        client = auth_data[user_id]['client']
        await client.sign_in(auth_data[user_id]['phone'], code)
        
        session = client.session.save()
        save_session_string(session)
        
        me = await client.get_me()
        await m.reply(f"✅ Авторизация успешна!\nАккаунт: {me.first_name}", reply_markup=get_main_keyboard(user_id))
        
        del auth_data[user_id]
        
    except SessionPasswordNeededError:
        auth_data[user_id]['state'] = 'waiting_password'
        await m.reply("🔐 Требуется пароль двухфакторки. Введи пароль:")
    except Exception as e:
        await m.reply(f"❌ Ошибка: {e}")

@dp.message_handler(lambda m: m.from_user.id in auth_data and auth_data[m.from_user.id]['state'] == 'waiting_password')
async def process_password(m: types.Message):
    user_id = m.from_user.id
    password = m.text.strip()
    
    try:
        client = auth_data[user_id]['client']
        await client.sign_in(password=password)
        
        session = client.session.save()
        save_session_string(session)
        
        me = await client.get_me()
        await m.reply(f"✅ Авторизация успешна!\nАккаунт: {me.first_name}", reply_markup=get_main_keyboard(user_id))
        
        del auth_data[user_id]
    except Exception as e:
        await m.reply(f"❌ Неверный пароль. Попробуй еще раз:")

# ========== ОБРАБОТКА ТЕКСТА ==========
@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_channel')
async def process_channel(m: types.Message):
    user_id = m.from_user.id
    link = m.text.strip()
    
    link = normalize_channel_url(link)
    name = extract_channel_name(link)
    
    channels = load_channels()
    
    for ch in channels:
        if ch['url'] == link:
            await m.reply("❌ Канал уже есть")
            del user_data[user_id]
            return
    
    channels.append({'name': name, 'url': link})
    save_channels(channels)
    
    del user_data[user_id]
    await m.reply(f"✅ Канал добавлен!\n{name}\n{link}\nВсего: {len(channels)}", reply_markup=get_main_keyboard(user_id))

@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_keywords')
async def process_keywords(m: types.Message):
    user_id = m.from_user.id
    keywords = m.text.strip()
    
    user_data[user_id]['keywords'] = keywords
    user_data[user_id]['state'] = 'waiting_period'
    
    await m.reply("⏱ За какой период?", reply_markup=get_period_keyboard())

@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_period')
async def process_period(m: types.Message):
    user_id = m.from_user.id
    text = m.text
    
    hours = 0
    
    if text == "🕐 1 час":
        hours = 1
    elif text == "🕒 3 часа":
        hours = 3
    elif text == "🕖 7 часов":
        hours = 7
    elif text == "📅 24 часа":
        hours = 24
    elif text == "📆 3 дня":
        hours = 72
    elif text == "📆 7 дней":
        hours = 168
    elif text == "⌨️ Свой":
        await m.reply("⌨️ Введи период (например: 30 минут, 2 часа, 5 дней):", reply_markup=get_custom_period_keyboard())
        user_data[user_id]['state'] = 'waiting_custom_period'
        return
    elif text == "◀️ Назад":
        await back(m)
        return
    else:
        await m.reply("❌ Выбери из кнопок")
        return
    
    user_data[user_id]['period_hours'] = hours
    user_data[user_id]['period_text'] = text
    
    await m.reply("🌍 Выбери язык перевода:", reply_markup=get_translation_keyboard())
    user_data[user_id]['state'] = 'waiting_translation'

@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_custom_period')
async def process_custom_period(m: types.Message):
    user_id = m.from_user.id
    text = m.text
    
    if text == "◀️ Назад":
        await m.reply("⏱ За какой период?", reply_markup=get_period_keyboard())
        user_data[user_id]['state'] = 'waiting_period'
        return
    
    hours = parse_period(text)
    
    if not hours or hours <= 0:
        await m.reply("❌ Не понимаю. Попробуй еще раз (30 минут, 2 часа, 5 дней):")
        return
    
    if hours > MAX_PERIOD_HOURS:
        await m.reply(f"❌ Максимум {MAX_PERIOD_DAYS} дней. Введи меньше:")
        return
    
    user_data[user_id]['period_hours'] = hours
    user_data[user_id]['period_text'] = text
    
    await m.reply("🌍 Выбери язык перевода:", reply_markup=get_translation_keyboard())
    user_data[user_id]['state'] = 'waiting_translation'

@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_translation')
async def process_translation(m: types.Message):
    user_id = m.from_user.id
    text = m.text
    
    if text == "◀️ Назад":
        await m.reply("⏱ За какой период?", reply_markup=get_period_keyboard())
        user_data[user_id]['state'] = 'waiting_period'
        return
    
    lang_map = {
        "🇷🇺 На русский": 'ru',
        "🇬🇧 На английский": 'en',
        "🇩🇪 На немецкий": 'de',
        "🇫🇷 На французский": 'fr',
        "🇪🇸 На испанский": 'es',
        "🔄 Без перевода": None
    }
    
    if text in lang_map:
        user_data[user_id]['translation'] = lang_map[text]
        await m.reply("🖼️ Выбери формат:", reply_markup=get_image_keyboard())
        user_data[user_id]['state'] = 'waiting_image'
    else:
        await m.reply("❌ Выбери из кнопок")

@dp.message_handler(lambda m: m.from_user.id in user_data and user_data[m.from_user.id].get('state') == 'waiting_image')
async def process_image(m: types.Message):
    user_id = m.from_user.id
    text = m.text
    
    if text == "◀️ Назад":
        await m.reply("🌍 Выбери язык перевода:", reply_markup=get_translation_keyboard())
        user_data[user_id]['state'] = 'waiting_translation'
        return
    
    save_images = (text == "🖼️ С картинками")
    user_data[user_id]['save_images'] = save_images
    
    search_type = user_data[user_id]['search_type']
    
    await m.reply(f"🔍 Начинаю поиск...\nЭто может занять время", reply_markup=get_main_keyboard(user_id))
    
    if search_type == 'channels':
        await collect_from_channels(user_id)
    elif search_type == 'global':
        await collect_global(user_id)
    elif search_type == 'web':
        await collect_web(user_id)
    elif search_type == 'combined':
        await collect_combined(user_id)

# ========== СБОР ДАННЫХ ==========
async def collect_from_channels(user_id):
    client = None
    try:
        stop_flags[user_id] = False
        
        data = user_data[user_id]
        keywords = data['keywords']
        channels = data['channels']
        hours = data['period_hours']
        save_images = data['save_images']
        trans_lang = data.get('translation')
        
        await bot.send_message(user_id, "🔄 Подключаюсь...")
        
        client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
        await client.connect()
        
        if not await client.is_user_authorized():
            await bot.send_message(user_id, "❌ Сессия не активна")
            return
        
        doc = Document()
        doc.add_heading('Поиск по каналам', 0).alignment = 1
        doc.add_paragraph(f"Дата: {datetime.now(UTC_PLUS_10).strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Слова: {keywords}")
        doc.add_paragraph(f"Период: {data['period_text']}")
        doc.add_paragraph()
        
        total = 0
        processed = 0
        channels_with = 0
        
        now = datetime.now().astimezone()
        start = now - timedelta(hours=hours)
        
        for channel in channels:
            if stop_flags.get(user_id):
                break
            
            processed += 1
            await bot.send_message(user_id, f"📱 {processed}/{len(channels)} {channel['name']}")
            
            try:
                url = normalize_channel_url(channel['url'])
                entity = await client.get_entity(url)
                
                posts = []
                count = 0
                
                async for msg in client.iter_messages(entity, offset_date=now):
                    if count % 10 == 0 and stop_flags.get(user_id):
                        break
                    
                    if msg.date:
                        d = msg.date
                        if d.tzinfo is None:
                            d = d.replace(tzinfo=timezone.utc)
                        if d < start:
                            break
                    
                    if keywords.lower() == 'все' or (msg.text and keywords.lower() in msg.text.lower()):
                        posts.append(msg)
                        count += 1
                        total += 1
                
                if posts:
                    channels_with += 1
                    doc.add_heading(f"Канал: {channel['name']}", level=1)
                    doc.add_paragraph(f"Ссылка: {url}")
                    doc.add_paragraph()
                    
                    for msg in posts:
                        p = doc.add_paragraph()
                        p.add_run(f"📅 {format_datetime_utc10(msg.date)}\n").bold = True
                        
                        if msg.text:
                            text = msg.text[:3000]
                            if trans_lang:
                                trans = await translator.translate(text, trans_lang)
                                if trans['was_translated']:
                                    p.add_run(f"[Переведено с {trans['src_name']}]\n\n{trans['translated']}")
                                else:
                                    p.add_run(text)
                            else:
                                p.add_run(text)
                        
                        if save_images and msg.media:
                            path = await download_media(msg, user_id, client)
                            if path:
                                doc.add_paragraph()
                                add_image_to_doc(doc, path)
                        
                        if msg.id:
                            username = url.split('/')[-1]
                            link = f"https://t.me/{username}/{msg.id}"
                            p = doc.add_paragraph()
                            add_hyperlink(p, "🔗 Ссылка", link)
                        
                        doc.add_paragraph()
                    
                    doc.add_paragraph(f"✅ Найдено: {count}")
                    doc.add_page_break()
                
            except Exception as e:
                await bot.send_message(user_id, f"⚠️ Ошибка: {str(e)[:100]}")
        
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f"Обработано: {processed}/{len(channels)}")
        doc.add_paragraph(f"С постами: {channels_with}")
        doc.add_paragraph(f"Всего постов: {total}")
        
        if total == 0:
            await bot.send_message(user_id, "📭 Ничего не найдено")
            return
        
        filename = f"report_{user_id}_{int(time.time())}.docx"
        doc.save(filename)
        
        with open(filename, 'rb') as f:
            await bot.send_document(user_id, f, caption=f"✅ Готово! Найдено: {total}")
        
        os.remove(filename)
        cleanup_temp_files(user_id)
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {e}")
    finally:
        if client:
            await client.disconnect()
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

async def collect_global(user_id):
    """Глобальный поиск по всему Telegram (исправлено)"""
    client = None
    try:
        stop_flags[user_id] = False
        
        data = user_data[user_id]
        keywords = data['keywords']
        hours = data['period_hours']
        save_images = data['save_images']
        trans_lang = data.get('translation')
        
        await bot.send_message(user_id, "🔄 Подключаюсь...")
        
        client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
        await client.connect()
        
        if not await client.is_user_authorized():
            await bot.send_message(user_id, "❌ Сессия не активна")
            return
        
        await bot.send_message(user_id, "🌍 Ищу по всему Telegram...")
        
        doc = Document()
        doc.add_heading('Глобальный поиск в Telegram', 0).alignment = 1
        doc.add_paragraph(f"Дата: {datetime.now(UTC_PLUS_10).strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Слова: {keywords}")
        doc.add_paragraph(f"Период: {data['period_text']}")
        doc.add_paragraph()
        
        start = datetime.now().astimezone() - timedelta(hours=hours)
        
        try:
            result = await client(SearchGlobalRequest(
                q=keywords,
                filter=InputMessagesFilterEmpty(),
                min_date=start,
                max_date=datetime.now().astimezone(),
                offset_rate=0,
                offset_peer=None,
                offset_id=0,
                limit=50
            ))
            
            total = 0
            
            if hasattr(result, 'messages') and result.messages:
                for msg in result.messages:
                    if stop_flags.get(user_id):
                        break
                    
                    if not hasattr(msg, 'message') or not msg.message:
                        continue
                    
                    total += 1
                    
                    # Получаем информацию о чате с проверкой
                    chat_title = 'Неизвестный чат'
                    chat_username = None
                    
                    try:
                        if msg.peer_id:
                            chat = await client.get_entity(msg.peer_id)
                            chat_title = getattr(chat, 'title', 'Неизвестный чат')
                            chat_username = getattr(chat, 'username', None)
                    except:
                        pass
                    
                    doc.add_heading(f"Чат: {chat_title}", level=2)
                    if chat_username:
                        doc.add_paragraph(f"https://t.me/{chat_username}")
                    
                    p = doc.add_paragraph()
                    p.add_run(f"📅 {format_datetime_utc10(msg.date)}\n").bold = True
                    
                    if msg.message:
                        text = msg.message[:2000]
                        if trans_lang:
                            trans = await translator.translate(text, trans_lang)
                            if trans['was_translated']:
                                p.add_run(f"[Переведено с {trans['src_name']}]\n\n{trans['translated']}")
                            else:
                                p.add_run(text)
                        else:
                            p.add_run(text)
                    
                    if chat_username and msg.id:
                        link = f"https://t.me/{chat_username}/{msg.id}"
                        p = doc.add_paragraph()
                        add_hyperlink(p, "🔗 Ссылка", link)
                    
                    doc.add_paragraph()
            
            doc.add_heading('Статистика', level=1)
            doc.add_paragraph(f"Всего: {total}")
            
            if total == 0:
                await bot.send_message(user_id, "📭 Ничего не найдено")
                return
            
            filename = f"global_{user_id}_{int(time.time())}.docx"
            doc.save(filename)
            
            with open(filename, 'rb') as f:
                await bot.send_document(user_id, f, caption=f"✅ Готово! Найдено: {total}")
            
            os.remove(filename)
            
        except Exception as e:
            await bot.send_message(user_id, f"❌ Ошибка поиска: {e}")
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {e}")
    finally:
        if client:
            await client.disconnect()
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

async def collect_web(user_id):
    try:
        stop_flags[user_id] = False
        
        data = user_data[user_id]
        keywords = data['keywords']
        web_type = data.get('web_type', 'all')
        trans_lang = data.get('translation')
        
        await bot.send_message(user_id, "🌐 Ищу в интернете...")
        
        doc = Document()
        doc.add_heading('Поиск в интернете', 0).alignment = 1
        doc.add_paragraph(f"Дата: {datetime.now(UTC_PLUS_10).strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Слова: {keywords}")
        doc.add_paragraph()
        
        total = 0
        
        if web_type == 'all' or web_type == 'google':
            results = await web_search.search_google(keywords, 10)
            if results:
                doc.add_heading('Google', level=1)
                for r in results:
                    if stop_flags.get(user_id):
                        break
                    doc.add_heading(r['title'], level=2)
                    doc.add_paragraph(f"Источник: {r['source']}")
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, r['link'], r['link'])
                    if r['description']:
                        if trans_lang:
                            trans = await translator.translate(r['description'], trans_lang)
                            if trans['was_translated']:
                                doc.add_paragraph(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                            else:
                                doc.add_paragraph(r['description'])
                        else:
                            doc.add_paragraph(r['description'])
                    doc.add_paragraph()
                    total += 1
                doc.add_page_break()
        
        if web_type == 'all' or web_type == 'bing':
            results = await web_search.search_bing(keywords, 10)
            if results:
                doc.add_heading('Bing', level=1)
                for r in results:
                    if stop_flags.get(user_id):
                        break
                    doc.add_heading(r['title'], level=2)
                    doc.add_paragraph(f"Источник: {r['source']}")
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, r['link'], r['link'])
                    if r['description']:
                        if trans_lang:
                            trans = await translator.translate(r['description'], trans_lang)
                            if trans['was_translated']:
                                doc.add_paragraph(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                            else:
                                doc.add_paragraph(r['description'])
                        else:
                            doc.add_paragraph(r['description'])
                    doc.add_paragraph()
                    total += 1
                doc.add_page_break()
        
        if web_type == 'all' or web_type == 'yandex':
            results = await web_search.search_yandex(keywords, 10)
            if results:
                doc.add_heading('Yandex', level=1)
                for r in results:
                    if stop_flags.get(user_id):
                        break
                    doc.add_heading(r['title'], level=2)
                    doc.add_paragraph(f"Источник: {r['source']}")
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, r['link'], r['link'])
                    if r['description']:
                        if trans_lang:
                            trans = await translator.translate(r['description'], trans_lang)
                            if trans['was_translated']:
                                doc.add_paragraph(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                            else:
                                doc.add_paragraph(r['description'])
                        else:
                            doc.add_paragraph(r['description'])
                    doc.add_paragraph()
                    total += 1
                doc.add_page_break()
        
        if web_type == 'all' or web_type == 'duckduckgo':
            results = await web_search.search_duckduckgo(keywords, 10)
            if results:
                doc.add_heading('DuckDuckGo', level=1)
                for r in results:
                    if stop_flags.get(user_id):
                        break
                    doc.add_heading(r['title'], level=2)
                    doc.add_paragraph(f"Источник: {r['source']}")
                    p = doc.add_paragraph()
                    p.add_run("Ссылка: ").bold = True
                    add_hyperlink(p, r['link'], r['link'])
                    if r['description']:
                        if trans_lang:
                            trans = await translator.translate(r['description'], trans_lang)
                            if trans['was_translated']:
                                doc.add_paragraph(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                            else:
                                doc.add_paragraph(r['description'])
                        else:
                            doc.add_paragraph(r['description'])
                    doc.add_paragraph()
                    total += 1
                doc.add_page_break()
        
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f"Всего: {total}")
        
        if total == 0:
            await bot.send_message(user_id, "📭 Ничего не найдено")
            return
        
        filename = f"web_{user_id}_{int(time.time())}.docx"
        doc.save(filename)
        
        with open(filename, 'rb') as f:
            await bot.send_document(user_id, f, caption=f"✅ Готово! Найдено: {total}")
        
        os.remove(filename)
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {e}")
    finally:
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

async def collect_combined(user_id):
    try:
        stop_flags[user_id] = False
        
        data = user_data[user_id]
        keywords = data['keywords']
        channels = data.get('channels', [])
        hours = data['period_hours']
        save_images = data['save_images']
        trans_lang = data.get('translation')
        
        await bot.send_message(user_id, "⚡ Комбинированный поиск...")
        
        doc = Document()
        doc.add_heading('Комбинированный поиск', 0).alignment = 1
        doc.add_paragraph(f"Дата: {datetime.now(UTC_PLUS_10).strftime('%d.%m.%Y %H:%M')} (UTC+10)")
        doc.add_paragraph(f"Слова: {keywords}")
        doc.add_paragraph(f"Период: {data['period_text']}")
        doc.add_paragraph()
        
        total = 0
        
        # Интернет
        await bot.send_message(user_id, "🌐 Ищу в интернете...")
        doc.add_heading('ИНТЕРНЕТ', level=1)
        
        results = await web_search.search_google(keywords, 5)
        results += await web_search.search_bing(keywords, 5)
        
        for r in results[:15]:
            if stop_flags.get(user_id):
                break
            doc.add_heading(r['title'], level=2)
            doc.add_paragraph(f"Источник: {r['source']}")
            p = doc.add_paragraph()
            p.add_run("Ссылка: ").bold = True
            add_hyperlink(p, r['link'], r['link'])
            if r['description']:
                if trans_lang:
                    trans = await translator.translate(r['description'], trans_lang)
                    if trans['was_translated']:
                        doc.add_paragraph(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                    else:
                        doc.add_paragraph(r['description'])
                else:
                    doc.add_paragraph(r['description'])
            doc.add_paragraph()
            total += 1
        
        doc.add_page_break()
        
        # Глобальный Telegram
        if MASTER_SESSION and not stop_flags.get(user_id):
            await bot.send_message(user_id, "🌍 Ищу в Telegram...")
            doc.add_heading('TELEGRAM', level=1)
            
            client = TelegramClient(StringSession(MASTER_SESSION), API_ID, API_HASH)
            await client.connect()
            
            if await client.is_user_authorized():
                start = datetime.now().astimezone() - timedelta(hours=hours)
                
                try:
                    result = await client(SearchGlobalRequest(
                        q=keywords,
                        filter=InputMessagesFilterEmpty(),
                        min_date=start,
                        max_date=datetime.now().astimezone(),
                        offset_rate=0,
                        offset_peer=None,
                        offset_id=0,
                        limit=20
                    ))
                    
                    if hasattr(result, 'messages') and result.messages:
                        for msg in result.messages[:10]:
                            if stop_flags.get(user_id) or not hasattr(msg, 'message'):
                                break
                            
                            chat_title = 'Неизвестный чат'
                            chat_username = None
                            
                            try:
                                if msg.peer_id:
                                    chat = await client.get_entity(msg.peer_id)
                                    chat_title = getattr(chat, 'title', 'Неизвестный чат')
                                    chat_username = getattr(chat, 'username', None)
                            except:
                                pass
                            
                            doc.add_heading(f"Чат: {chat_title}", level=2)
                            if chat_username:
                                doc.add_paragraph(f"https://t.me/{chat_username}")
                            
                            p = doc.add_paragraph()
                            p.add_run(f"📅 {format_datetime_utc10(msg.date)}\n").bold = True
                            
                            if msg.message:
                                text = msg.message[:1000]
                                if trans_lang:
                                    trans = await translator.translate(text, trans_lang)
                                    if trans['was_translated']:
                                        p.add_run(f"[Переведено с {trans['src_name']}]\n{trans['translated']}")
                                    else:
                                        p.add_run(text)
                                else:
                                    p.add_run(text)
                            
                            if chat_username and msg.id:
                                link = f"https://t.me/{chat_username}/{msg.id}"
                                p = doc.add_paragraph()
                                add_hyperlink(p, "🔗 Ссылка", link)
                            
                            doc.add_paragraph()
                            total += 1
                except:
                    pass
            
            await client.disconnect()
            doc.add_page_break()
        
        doc.add_heading('Статистика', level=1)
        doc.add_paragraph(f"Всего: {total}")
        
        if total == 0:
            await bot.send_message(user_id, "📭 Ничего не найдено")
            return
        
        filename = f"combined_{user_id}_{int(time.time())}.docx"
        doc.save(filename)
        
        with open(filename, 'rb') as f:
            await bot.send_document(user_id, f, caption=f"✅ Готово! Найдено: {total}")
        
        os.remove(filename)
        
    except Exception as e:
        await bot.send_message(user_id, f"❌ Ошибка: {e}")
    finally:
        if user_id in user_data:
            del user_data[user_id]
        if user_id in stop_flags:
            del stop_flags[user_id]

# ========== КОЛБЭКИ ==========
@dp.callback_query_handler(lambda c: c.data.startswith('del_'))
async def delete_callback(call):
    await bot.answer_callback_query(call.id)
    
    user_id = call.from_user.id
    if not is_admin(user_id):
        await bot.send_message(user_id, "❌ Только для админа")
        return
    
    data = call.data
    channels = load_channels()
    
    if data == "del_all":
        save_channels([])
        await bot.send_message(user_id, "🗑 Все каналы удалены")
        return
    
    idx = int(data.split('_')[1])
    if idx < len(channels):
        deleted = channels.pop(idx)
        save_channels(channels)
        await bot.send_message(user_id, f"🗑 Удален: {deleted['name']}")

@dp.callback_query_handler(lambda c: c.data == 'back_to_main')
async def back_callback(call):
    await bot.answer_callback_query(call.id)
    user_id = call.from_user.id
    if user_id in user_data:
        del user_data[user_id]
    await bot.send_message(user_id, "Главное меню", reply_markup=get_main_keyboard(user_id))

# ========== НЕИЗВЕСТНЫЕ КОМАНДЫ ==========
@dp.message_handler()
async def unknown(m: types.Message):
    user_id = m.from_user.id
    
    if is_admin(user_id) and user_id in auth_data:
        state = auth_data[user_id].get('state')
        if state == 'waiting_phone':
            await process_phone(m)
        elif state == 'waiting_code':
            await process_code(m)
        elif state == 'waiting_password':
            await process_password(m)
        return
    
    if m.text and ('t.me/' in m.text or '@' in m.text):
        user_data[user_id] = {'state': 'waiting_channel'}
        await process_channel(m)
    else:
        await m.reply("Используй кнопки 👇", reply_markup=get_main_keyboard(user_id))

# ========== ЗАПУСК ==========
if __name__ == '__main__':
    MASTER_SESSION = load_master_session()
    channels = load_channels()
    
    print("=" * 40)
    print("🤖 Бот запущен")
    print(f"👑 Админ: {ADMIN_ID}")
    print(f"📊 Каналов: {len(channels)}")
    print(f"🌍 Перевод: есть")
    print("=" * 40)
    
    executor.start_polling(dp, skip_updates=True)
